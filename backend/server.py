import os
os.environ["OPENBLAS_NUM_THREADS"] = "1"
os.environ["MKL_NUM_THREADS"] = "1" 
os.environ["NUMEXPR_NUM_THREADS"] = "1"
os.environ["OMP_NUM_THREADS"] = "1"
os.environ["TOKENIZERS_PARALLELISM"] = "false"

from flask import Flask, request, jsonify, send_file, g, Blueprint
from flask_cors import CORS
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from werkzeug.security import generate_password_hash, check_password_hash
from flask_sqlalchemy import SQLAlchemy
from dotenv import load_dotenv
from flask_jwt_extended import jwt_required, get_jwt_identity, get_jwt
import PyPDF2
import pdfplumber
import io
from datetime import datetime, timezone, timedelta
import pytz
import uuid
import asyncio
import time
from functools import wraps
from asgiref.wsgi import WsgiToAsgi
import pymysql
from openai import OpenAI
import json
import traceback

# ============================================================
# LAZY LOADING - Heavy ML libraries loaded on demand only
# ============================================================
# REMOVED these direct imports that cause crashes on shared hosting:
# import chromadb
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langchain_openai import OpenAIEmbeddings
# from langchain_community.embeddings import HuggingFaceEmbeddings
# from sentence_transformers import SentenceTransformer

from services.extract import extract_text_from_pdf
from services.chunk import chunk_text
from services.vector import add_chunks, search

# PDF generation
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("ReportLab not available - PDF export disabled")


def log_notification_error(err):
    try:
        with open("notifications-error.log", "a") as f:
            f.write("\n=========================\n")
            f.write(str(err) + "\n")
            f.write(traceback.format_exc() + "\n")
            f.write("=========================\n")
    except:
        pass


# Import for multi-format support
try:
    from PIL import Image
except ImportError:
    Image = None


# ============================================================
# SIMPLE TEXT SPLITTER (Replaces langchain RecursiveCharacterTextSplitter)
# ============================================================
class SimpleTextSplitter:
    """Simple text splitter - no heavy dependencies"""
    def __init__(self, chunk_size=1000, chunk_overlap=200, length_function=len):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.length_function = length_function
    
    def split_text(self, text):
        if not text:
            return []
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            # Try to break at sentence/paragraph boundary
            if end < len(text):
                for sep in ['\n\n', '.\n', '. ', '\n', ' ']:
                    last_sep = text[start:end].rfind(sep)
                    if last_sep > self.chunk_size * 0.5:
                        end = start + last_sep + len(sep)
                        break
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start = end - self.chunk_overlap if end < len(text) else end
        return chunks


# This replaces your old text_splitter
text_splitter = SimpleTextSplitter(
    chunk_size=1000,
    chunk_overlap=200,
    length_function=len,
)


# ============================================================
# LAZY LOADER FOR HEAVY LIBRARIES
# ============================================================
class LazyLoader:
    """Load heavy ML libraries only when needed"""
    _cache = {}
    
    @classmethod
    def get_chromadb(cls):
        if 'chromadb' not in cls._cache:
            try:
                import chromadb
                cls._cache['chromadb'] = chromadb.Client()
            except Exception as e:
                print(f"⚠️ ChromaDB not available: {e}")
                cls._cache['chromadb'] = None
        return cls._cache['chromadb']
    
    @classmethod
    def get_sentence_transformer(cls):
        if 'sentence_transformer' not in cls._cache:
            try:
                from sentence_transformers import SentenceTransformer
                cls._cache['sentence_transformer'] = SentenceTransformer('all-MiniLM-L6-v2')
            except Exception as e:
                print(f"⚠️ SentenceTransformer not available: {e}")
                cls._cache['sentence_transformer'] = None
        return cls._cache['sentence_transformer']


# ============================================================
# FAKE VECTOR DB (Fallback when ChromaDB unavailable)
# ============================================================
class FakeVectorDB:
    def add(self, *args, **kwargs): pass
    def get_collection(self, *args, **kwargs):
        class C:
            def query(self, *args, **kwargs): return {"documents": [[]]}
        return C()

chroma_client = FakeVectorDB()
embeddings = None
print("⚠️ Vector Search Disabled (Shared Hosting Mode)")


# ==================== CORE FUNCTIONS ====================

def get_global_settings():
    """Get global settings for the application"""
    try:
        settings = GlobalSettings.query.first()
        if settings:
            return {
                'free_uploads_per_day': settings.free_uploads_per_day,
                'free_chats_per_day': settings.free_chats_per_day,
                'daily_token_cap': settings.daily_token_cap,
                'streak_bonus_base': settings.streak_bonus_base
            }
        else:
            return {
                'free_uploads_per_day': 5,
                'free_chats_per_day': 10,
                'daily_token_cap': 200,
                'streak_bonus_base': 10
            }
    except Exception as e:
        print(f"Error fetching global settings: {e}")
        return {
            'free_uploads_per_day': 5,
            'free_chats_per_day': 10,
            'daily_token_cap': 200,
            'streak_bonus_base': 10
        }


def get_stats(user_id):
    """Get user stats for token and subscription checking"""
    try:
        stats = UserStats.query.filter_by(user_id=user_id).first()
        if not stats:
            # Create default stats if not exist
            ist = pytz.timezone('Asia/Kolkata')
            now = datetime.now(ist)
            today = now.date()
            current_month_date = now.replace(day=1).date()

            stats = UserStats(
                user_id=user_id,
                tokens=100,
                monthly_tokens=100,
                xp=0,
                streak=1,
                level=1,
                last_active_date=today,
                last_reset=today,
                monthly_reset_date=current_month_date
            )
            db.session.add(stats)
            db.session.commit()
            print(f"✅ Created new user stats for user {user_id}")
        return stats
    except Exception as e:
        print(f"Error getting user stats: {e}")
        import traceback
        print(traceback.format_exc())
        return None


def spend_tokens(user_id, amount):
    """Spend tokens from user's balance"""
    try:
        stats = get_stats(user_id)
        if not stats or stats.tokens < amount:
            return False
        stats.tokens -= amount
        db.session.commit()
        return True
    except Exception as e:
        print(f"Error spending tokens: {e}")
        db.session.rollback()
        return False


def add_tokens(user_id, amount):
    """Add tokens to user's balance"""
    try:
        stats = get_stats(user_id)
        if not stats:
            return False
        stats.tokens += amount
        db.session.commit()
        return True
    except Exception as e:
        print(f"Error adding tokens: {e}")
        db.session.rollback()
        return False


def check_monthly_token_reset(user_id):
    """Check if monthly tokens need to be reset and handle it"""
    try:
        stats = get_stats(user_id)
        if not stats:
            return False

        ist = pytz.timezone('Asia/Kolkata')
        now = datetime.now(ist)
        current_month_date = now.replace(day=1).date()

        # Convert monthly_reset_date to date if needed
        last_reset = stats.monthly_reset_date
        if last_reset is not None and hasattr(last_reset, 'date') and callable(getattr(last_reset, 'date')):
            last_reset = last_reset.date()

        # If no monthly reset date set or it's a new month, reset tokens
        if last_reset is None or last_reset < current_month_date:
            is_first_month = last_reset is None
            base_tokens = 100
            bonus_tokens = 100 if is_first_month else 0
            stats.monthly_tokens = base_tokens + bonus_tokens
            stats.monthly_reset_date = current_month_date

            if is_first_month:
                stats.tokens += stats.monthly_tokens

            db.session.commit()
            return True

        return False
    except Exception as e:
        print(f"Error checking monthly token reset: {e}")
        try:
            db.session.rollback()
        except:
            pass
        return False


def get_subscription_status(user_id):
    """Get user's subscription status"""
    try:
        subscription = Subscription.query.filter_by(user_id=user_id, active=True).first()
        if subscription:
            ist = pytz.timezone('Asia/Kolkata')
            if subscription.expires_at and subscription.expires_at < datetime.now(ist):
                subscription.active = False
                db.session.commit()
                return {'tier': 'free', 'active': True, 'expires_at': None}
            return {
                'tier': subscription.tier,
                'active': True,
                'expires_at': subscription.expires_at.isoformat() if subscription.expires_at else None
            }
        return {'tier': 'free', 'active': True, 'expires_at': None}
    except Exception as e:
        print(f"Error getting subscription status: {e}")
        return {'tier': 'free', 'active': True, 'expires_at': None}


def update_streak_and_xp_fast(user_id, activity_type="login"):
    """Fast version of streak/XP update for login"""
    try:
        stats = UserStats.query.filter_by(user_id=user_id).first()
        if not stats:
            return False

        ist = pytz.timezone('Asia/Kolkata')
        now = datetime.now(ist)
        today = now.date()

        # Streak update
        if stats.last_active_date is None:
            stats.streak = 1
        elif stats.last_active_date != today:
            if stats.last_active_date == today - timedelta(days=1):
                stats.streak += 1
            else:
                stats.streak = 1

        stats.last_active_date = today

        # XP update
        xp_gains = {
            'login': 5, 'upload': 15, 'flashcard_generation': 20,
            'flashcard_review': 5, 'chat': 10, 'file_chat': 10
        }
        stats.xp += xp_gains.get(activity_type, 5)

        # Level calculation
        import math
        new_level = max(1, math.floor(stats.xp / 100) + 1)
        if new_level > stats.level:
            stats.level = new_level
            stats.tokens += 25  # Level up reward

        return True
    except Exception as e:
        print(f"Error in fast streak/XP update: {e}")
        try:
            db.session.rollback()
        except:
            pass
        return False


def get_user_token_info(user_id):
    """Get comprehensive token information for user"""
    try:
        check_monthly_token_reset(user_id)
        stats = get_stats(user_id)
        subscription = get_subscription_status(user_id)

        if not stats:
            return None

        ist = pytz.timezone('Asia/Kolkata')
        now = datetime.now(ist)

        # Calculate next monthly reset
        if stats.monthly_reset_date:
            if hasattr(stats.monthly_reset_date, 'date') and callable(getattr(stats.monthly_reset_date, 'date')):
                reset_date = stats.monthly_reset_date
            else:
                reset_date = datetime.combine(stats.monthly_reset_date, datetime.min.time())
                reset_date = ist.localize(reset_date)
            next_reset = (reset_date + timedelta(days=32)).replace(day=1)
        else:
            next_reset = (now + timedelta(days=32)).replace(day=1)

        # Days until reset
        if hasattr(next_reset, 'date'):
            days_until_reset = (next_reset.date() - now.date()).days
        else:
            days_until_reset = 30

        return {
            'current_tokens': stats.tokens,
            'monthly_tokens_remaining': stats.monthly_tokens,
            'subscription': subscription,
            'streak_days': stats.streak,
            'level': stats.level,
            'xp': stats.xp,
            'next_monthly_reset': next_reset.isoformat() if hasattr(next_reset, 'isoformat') else str(next_reset),
            'days_until_reset': max(0, days_until_reset),
            'xp_to_next_level': max(0, ((stats.level + 1) ** 2 * 10) - stats.xp)
        }
    except Exception as e:
        print(f"Error getting user token info: {e}")
        import traceback
        print(traceback.format_exc())
        return None


def track_event(event_type, event_data=None):
    """Track analytics events"""
    try:
        event = Analytics(
            event_id=str(uuid.uuid4()),
            user_id=event_data.get('user_id') if event_data else None,
            event_type=event_type,
            event_data=event_data,
            timestamp=datetime.now(pytz.timezone('Asia/Kolkata'))
        )
        db.session.add(event)
        db.session.commit()
    except Exception as e:
        print(f"Error tracking event: {e}")
        try:
            db.session.rollback()
        except:
            pass


def can_chat(user_id):
    """Check if user can chat based on limits and tokens"""
    try:
        check_monthly_token_reset(user_id)
        stats = get_stats(user_id)
        if not stats:
            return False

        subscription = get_subscription_status(user_id)
        if subscription and subscription.get('tier') == 'premium':
            return True

        # Check daily free chat limit
        today = date.today()
        chats_today = ChatLog.query.filter(
            ChatLog.user_id == user_id,
            ChatLog.timestamp >= datetime.combine(today, datetime.min.time())
        ).count()

        settings = get_global_settings()
        free_limit = settings['free_chats_per_day']

        if chats_today < free_limit:
            return True

        # Use tokens if over free limit
        return stats.tokens >= 3
    except Exception as e:
        print(f"Error checking can_chat: {e}")
        return False

# Admin role decorator
def admin_required(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        from flask_jwt_extended import get_jwt
        claims = get_jwt()
        if claims.get("role") != "admin":
            return jsonify({"error": "Admin access required"}), 403
        return func(*args, **kwargs)
    return wrapper

# Admin role decorator
def admin_required(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        from flask_jwt_extended import get_jwt
        claims = get_jwt()
        if claims.get("role") != "admin":
            return jsonify({"error": "Admin access required"}), 403
        return func(*args, **kwargs)
    return wrapper


load_dotenv()

app = Flask(__name__)
app.config['JWT_SECRET_KEY'] = os.environ.get('JWT_SECRET_KEY', 'impify-secret-key-change-in-production')
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(days=7)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# MySQL Database Configuration
mysql_url = os.environ.get('MYSQL_URL', 'mysql+pymysql://visasyst:FLLq37d)s9B:d6@localhost:3306/visasyst_impify')
app.config['SQLALCHEMY_DATABASE_URI'] = mysql_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)
jwt = JWTManager(app)

# Create all database tables
with app.app_context():
    db.create_all()
    print("✅ Database tables created successfully")
CORS(app, supports_credentials=True, resources={
    r"/api/*": {
        "origins": [
            "https://impify.visasystem.in",
            "http://impify.visasystem.in", 
            "http://localhost:3000",
            "http://127.0.0.1:3000",
            "http://localhost:5000",
            "http://127.0.0.1:5000"
        ],
        "allow_headers": ["Content-Type", "Authorization", "X-Requested-With"],
        "expose_headers": ["Content-Type", "Authorization"],
        "methods": ["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
        "supports_credentials": True
    }
})

# @app.after_request
# def after_request(response):
#     response.headers.add("Access-Control-Allow-Origin", "http://localhost:3000")
#     response.headers.add("Access-Control-Allow-Headers", "Content-Type,Authorization")
#     response.headers.add("Access-Control-Allow-Methods", "GET,POST,PUT,PATCH,DELETE,OPTIONS")
#     return response

# SQLAlchemy Models
class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.String(36), primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    name = db.Column(db.String(100))
    password = db.Column(db.String(255), nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))
    
    def set_password(self, password):
        """Hash and set password"""
        from werkzeug.security import generate_password_hash
        self.password = generate_password_hash(password)
    
    def check_password(self, password):
        """Check password against hash"""
        from werkzeug.security import check_password_hash
        return check_password_hash(self.password, password)

class Note(db.Model):
    __tablename__ = 'notes'
    id = db.Column(db.String(36), primary_key=True)
    user_id = db.Column(db.String(36), db.ForeignKey('users.id'), nullable=False)
    title = db.Column(db.String(255), nullable=False)
    original_filename = db.Column(db.String(255))
    note_type = db.Column(db.String(50), default='general')
    content = db.Column(db.Text, nullable=False)
    file_size = db.Column(db.Integer)
    processing_time = db.Column(db.Float)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')), onupdate=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))

class Consent(db.Model):
    __tablename__ = 'consents'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    user_id = db.Column(db.String(36), db.ForeignKey('users.id'), nullable=False)
    consented = db.Column(db.Boolean, default=False)
    consent_text = db.Column(db.Text)
    consent_date = db.Column(db.DateTime)
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')), onupdate=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))

class Analytics(db.Model):
    __tablename__ = 'analytics'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    event_id = db.Column(db.String(36), unique=True, nullable=False)
    user_id = db.Column(db.String(36))
    event_type = db.Column(db.String(100), nullable=False)
    event_data = db.Column(db.JSON)
    timestamp = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))
    ip_address = db.Column(db.String(45))
    user_agent = db.Column(db.Text)
class SupportTicket(db.Model):
    __tablename__ = 'support_tickets'
    id = db.Column(db.String(36), primary_key=True)
    title = db.Column(db.String(255), nullable=False)
    description = db.Column(db.Text, nullable=False)
    category = db.Column(db.String(50), nullable=False)  # bug, feature, improvement, question, other
    priority = db.Column(db.String(20), default='medium')  # low, medium, high, urgent
    status = db.Column(db.String(20), default='open')  # open, in_progress, resolved, closed
    author_id = db.Column(db.String(36), nullable=False)  # Can be volunteer or regular user
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')), onupdate=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))
class CommunityPost(db.Model):
    __tablename__ = 'community_posts'
    id = db.Column(db.String(36), primary_key=True)
    author_id = db.Column(db.String(36), nullable=False)  # Regular user ID
    author_name = db.Column(db.String(100), nullable=False)  # Store name for display
    author_email = db.Column(db.String(120), nullable=False)  # Store email for display
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')), onupdate=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))

class CommunityLike(db.Model):
    __tablename__ = 'community_likes'
    id = db.Column(db.String(36), primary_key=True)
    post_id = db.Column(db.String(36), db.ForeignKey('community_posts.id'), nullable=False)
    user_id = db.Column(db.String(36), nullable=False)  # Regular user ID
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))

    post = db.relationship('CommunityPost', backref=db.backref('likes', lazy=True))

class Notification(db.Model):
    __tablename__ = 'notifications'
    id = db.Column(db.String(36), primary_key=True)
    user_id = db.Column(db.String(36), nullable=False)  # Target user ID
    title = db.Column(db.String(100), nullable=False)
    message = db.Column(db.Text, nullable=False)
    type = db.Column(db.String(20), default='info')  # info, success, warning, error
    is_read = db.Column(db.Boolean, default=False)
    is_broadcast = db.Column(db.Boolean, default=False)  # True if sent to all users
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))
    expires_at = db.Column(db.DateTime, nullable=True)
    created_by = db.Column(db.String(36), nullable=True)  # Admin user ID who created it

class PasswordResetToken(db.Model):
    __tablename__ = 'password_reset_tokens'
    id = db.Column(db.String(36), primary_key=True)
    user_id = db.Column(db.String(36), db.ForeignKey('users.id'), nullable=False)
    token = db.Column(db.String(255), unique=True, nullable=False)
    expires_at = db.Column(db.DateTime, nullable=False)
    used = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))

    user = db.relationship('User', backref=db.backref('reset_tokens', lazy=True))

class Flashcard(db.Model):
    __tablename__ = 'flashcards'
    id = db.Column(db.String(36), primary_key=True)
    user_id = db.Column(db.String(36), db.ForeignKey('users.id'), nullable=False)
    note_id = db.Column(db.String(36), db.ForeignKey('notes.id'), nullable=True)  # Can be standalone or linked to a note
    question = db.Column(db.Text, nullable=False)
    answer = db.Column(db.Text, nullable=False)
    difficulty_score = db.Column(db.Float, default=0.0)  # 0.0 (easy) to 1.0 (hard)
    last_reviewed = db.Column(db.DateTime, nullable=True)
    next_review = db.Column(db.DateTime, nullable=True)
    review_count = db.Column(db.Integer, default=0)
    correct_count = db.Column(db.Integer, default=0)  # Times answered correctly
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')), onupdate=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))

    # Relationship
    note = db.relationship('Note', backref=db.backref('flashcards', lazy=True))

class UserStats(db.Model):
    __tablename__ = "user_stats"
    user_id = db.Column(db.String(36), primary_key=True)
    tokens = db.Column(db.Integer, default=0)
    monthly_tokens = db.Column(db.Integer, default=100)
    xp = db.Column(db.Integer, default=0)
    streak = db.Column(db.Integer, default=0)
    level = db.Column(db.Integer, default=1)
    last_active_date = db.Column(db.Date)
    last_reset = db.Column(db.Date)
    monthly_reset_date = db.Column(db.Date)

class UserDailyUploads(db.Model):
    __tablename__ = "user_daily_uploads"
    id = db.Column(db.String(36), primary_key=True)
    user_id = db.Column(db.String(36))
    upload_date = db.Column(db.Date)
    uploads_count = db.Column(db.Integer, default=0)

class Subscription(db.Model):
    __tablename__ = "subscriptions"
    id = db.Column(db.String(36), primary_key=True)
    user_id = db.Column(db.String(36))
    tier = db.Column(db.String(20))
    active = db.Column(db.Boolean, default=True)
    expires_at = db.Column(db.DateTime)
    created_at = db.Column(db.DateTime, server_default=db.func.now())

class Referral(db.Model):
    __tablename__ = "referrals"
    id = db.Column(db.String(36), primary_key=True)
    referrer_id = db.Column(db.String(36))
    referee_id = db.Column(db.String(36))
    code = db.Column(db.String(50))
    bonus_given = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, server_default=db.func.now())

class GlobalSettings(db.Model):
    __tablename__ = "global_settings"
    id = db.Column(db.Integer, primary_key=True)
    free_uploads_per_day = db.Column(db.Integer, default=5)
    free_chats_per_day = db.Column(db.Integer, default=10)
    daily_token_cap = db.Column(db.Integer, default=200)
    streak_bonus_base = db.Column(db.Integer, default=10)


class ChatLog(db.Model):
    __tablename__ = "chat_logs"
    id = db.Column(db.String(36), primary_key=True)
    user_id = db.Column(db.String(36), db.ForeignKey('users.id'), nullable=False)
    message = db.Column(db.Text, nullable=False)
    response = db.Column(db.Text, nullable=False)
    tokens_used = db.Column(db.Integer, default=0)
    timestamp = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))

class TokenPack(db.Model):
    __tablename__ = "token_packs"
    id = db.Column(db.String(36), primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    tokens = db.Column(db.Integer, nullable=False)
    price_in_inr = db.Column(db.Integer, nullable=False)
    description = db.Column(db.Text)
    active = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))

class PurchaseHistory(db.Model):
    __tablename__ = "purchase_history"
    id = db.Column(db.String(36), primary_key=True)
    user_id = db.Column(db.String(36), db.ForeignKey('users.id'), nullable=False)
    pack_id = db.Column(db.String(36), db.ForeignKey('token_packs.id'), nullable=True)
    tokens_credited = db.Column(db.Integer, nullable=False)
    amount_paid_in_inr = db.Column(db.Integer, nullable=False)
    payment_id = db.Column(db.String(255), nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(pytz.timezone('Asia/Kolkata')))


# AI clients with environment keys
openai_client = None
if os.environ.get("OPENAI_API_KEY"):
    openai_client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

# Gemini client removed - not using Gemini anymore

# Pluggable LLM Client
class LLMClient:
    def __init__(self, provider="openai", use_premium=False):
        self.provider = provider
        self.use_premium = use_premium
        if provider == "openai":
            from openai import OpenAI
            self.client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
        elif provider == "ollama":
            import requests
            self.base = os.environ.get("OLLAMA_URL", "http://localhost:11434")
        # add gemini etc.

    def chat(self, system, user):
        if self.provider == "openai":
            # Use gpt-4o for premium users, gpt-4o-mini for free users
            model = "gpt-4o" if self.use_premium else os.environ.get("OPENAI_MODEL", "gpt-4o-mini")
            r = self.client.chat.completions.create(
                model=model,
                messages=[{"role":"system","content":system},{"role":"user","content":user}]
            )
            return r.choices[0].message.content
        elif self.provider == "ollama":
            import requests
            r = requests.post(f"{self.base}/api/chat", json={
                "model": os.environ.get("OLLAMA_MODEL","llama3:8b"),
                "messages":[{"role":"system","content":system},{"role":"user","content":user}]
            }, timeout=120)
            r.raise_for_status()
            return r.json()["message"]["content"]

# Data logging for training
def log_training_example(user_id, source_file, prompt, input_text, output_text, meta):
    rec = {
        "user_id": user_id,
        "source": source_file,
        "instruction": prompt,
        "input": input_text[:4000],
        "output": output_text,
        "meta": meta,
        "created_at": datetime.now(pytz.timezone('Asia/Kolkata')).isoformat()
    }
    os.makedirs("training", exist_ok=True)
    with open("training/log.jsonl", "a", encoding="utf-8") as f:
        f.write(json.dumps(rec, ensure_ascii=False) + "\n")

# Vector DB and Embeddings setup
# 🚫 Disable ChromaDB on shared hosting (fallback mode)
class FakeVectorDB:
    def add(self, *args, **kwargs): pass
    def get_collection(self, *args, **kwargs):
        class C:
            def query(self, *args, **kwargs): return {"documents": [[]]}
        return C()

chroma_client = FakeVectorDB()
embeddings = None

print("⚠️ Vector Search Disabled (Shared Hosting Mode)")


# ==================== ANALYTICS TRACKING ====================

def track_event(event_type, event_data=None):
    """Track analytics events"""
    try:
        user_id = g.get('user_id', 'anonymous')
        event = Analytics(
            event_id=str(uuid.uuid4()),
            user_id=user_id,
            event_type=event_type,
            event_data=event_data or {},
            timestamp=datetime.now(pytz.timezone('Asia/Kolkata')),
            ip_address=request.remote_addr,
            user_agent=request.headers.get('User-Agent', '')
        )
        db.session.add(event)
        db.session.commit()
    except Exception as e:
        print(f"Analytics tracking error: {e}")
        db.session.rollback()

def track_usage(func):
    """Decorator to track API usage and timing"""
    @wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.time()
        try:
            result = func(*args, **kwargs)
            duration = time.time() - start_time
            
            # Track successful usage
            track_event('api_call', {
                'endpoint': request.endpoint,
                'method': request.method,
                'duration_ms': round(duration * 1000, 2),
                'status': 'success'
            })
            return result
        except Exception as e:
            duration = time.time() - start_time
            track_event('api_call', {
                'endpoint': request.endpoint,
                'method': request.method,
                'duration_ms': round(duration * 1000, 2),
                'status': 'error',
                'error': str(e)
            })
            raise
    return wrapper

@app.before_request
def before_request():
    """Set user context for analytics"""
    try:
        from flask_jwt_extended import verify_jwt_in_request, get_jwt_identity
        verify_jwt_in_request(optional=True)
        g.user_id = get_jwt_identity()
    except:
        g.user_id = None

# ==================== CONSENT MANAGEMENT ====================

def has_user_consent(user_id):
    """Check if user has given consent for AI processing"""
    consent = Consent.query.filter_by(user_id=user_id).first()
    return consent.consented if consent else False

@app.route('/api/consent/status', methods=['GET'])
@jwt_required()
@track_usage
def get_consent_status():
    user_id = get_jwt_identity()
    consent = Consent.query.filter_by(user_id=user_id).first()

    if not consent:
        return jsonify({
            'consented': False,
            'consent_date': None,
            'consent_text': None
        }), 200

    return jsonify({
        'consented': consent.consented,
        'consent_date': consent.consent_date.isoformat() if consent.consent_date else None,
        'consent_text': consent.consent_text or ''
    }), 200

@app.route('/api/consent/update', methods=['POST'])
@jwt_required()
@track_usage
def update_consent():
    user_id = get_jwt_identity()
    data = request.get_json()
    consented = data.get('consented', False)
    consent_text = data.get('consent_text', '')

    consent = Consent.query.filter_by(user_id=user_id).first()
    if not consent:
        consent = Consent(user_id=user_id)

    consent.consented = consented
    consent.consent_text = consent_text
    consent.consent_date = datetime.now(pytz.timezone('Asia/Kolkata'))
    consent.updated_at = datetime.now(pytz.timezone('Asia/Kolkata'))

    db.session.add(consent)
    db.session.commit()

    # Track consent decision
    track_event('consent_updated', {
        'consented': consented,
        'action': 'opt_in' if consented else 'opt_out'
    })

    return jsonify({'message': 'Consent updated successfully', 'consented': consented}), 200

# ==================== RATE LIMITING ====================

# File chat configuration
MAX_CONTEXT_CHUNKS = 5           # how many chunks to inject
ESTIMATED_TOKEN_COST = 2         # optional token-charge per chat (use with your token system)
MAX_CHUNK_PREVIEW = 300          # chars to include in returned chunk previews

# ==================== SUBSCRIPTION PLANS ====================
SUBSCRIPTION_PLANS = {
    'free': {
        'id': 'free',
        'name': 'Free',
        'price': 0,
        'currency': 'INR',
        'period': 'forever',
        'features': [
            '100 tokens per month',
            '5 uploads per day',
            '10 chats per day',
            'Basic AI features',
            'Community access'
        ],
        'tokens_per_month': 100,
        'uploads_per_day': 5,
        'chats_per_day': 10
    },
    'basic': {
        'id': 'basic',
        'name': 'Basic',
        'price': 199,
        'currency': 'INR',
        'period': 'month',
        'features': [
            '500 tokens per month',
            '20 uploads per day',
            '50 chats per day',
            'Priority AI processing',
            'Email support'
        ],
        'tokens_per_month': 500,
        'uploads_per_day': 20,
        'chats_per_day': 50
    },
    'pro': {
        'id': 'pro',
        'name': 'Pro',
        'price': 499,
        'currency': 'INR',
        'period': 'month',
        'features': [
            '2000 tokens per month',
            'Unlimited uploads',
            'Unlimited chats',
            'Advanced AI features',
            'Priority support',
            'Early access to new features'
        ],
        'tokens_per_month': 2000,
        'uploads_per_day': -1,  # -1 = unlimited
        'chats_per_day': -1
    },
    'premium': {
        'id': 'premium',
        'name': 'Premium',
        'price': 999,
        'currency': 'INR',
        'period': 'month',
        'features': [
            'Unlimited tokens',
            'Unlimited uploads',
            'Unlimited chats',
            'GPT-4o access',
            'API access',
            '24/7 Priority support',
            'Custom integrations'
        ],
        'tokens_per_month': -1,  # -1 = unlimited
        'uploads_per_day': -1,
        'chats_per_day': -1
    }
}

def get_global_settings():
    """Get global settings from database, with defaults"""
    try:
        settings = GlobalSettings.query.first()
        if settings:
            return {
                'free_uploads_per_day': settings.free_uploads_per_day,
                'free_chats_per_day': settings.free_chats_per_day,
                'daily_token_cap': settings.daily_token_cap,
                'streak_bonus_base': settings.streak_bonus_base
            }
        else:
            # Return defaults if no settings exist
            return {
                'free_uploads_per_day': 5,
                'free_chats_per_day': 10,
                'daily_token_cap': 200,
                'streak_bonus_base': 10
            }
    except Exception as e:
        print(f"Error fetching global settings: {e}")
        # Return defaults on error
        return {
            'free_uploads_per_day': 5,
            'free_chats_per_day': 10,
            'daily_token_cap': 200,
            'streak_bonus_base': 10
        }


def spend_tokens(user_id, amount):
    """Spend tokens from user's balance"""
    try:
        stats = get_stats(user_id)
        if not stats or stats.tokens < amount:
            return False

        stats.tokens -= amount
        db.session.commit()
        return True
    except Exception as e:
        print(f"Error spending tokens: {e}")
        db.session.rollback()
        return False

def add_tokens(user_id, amount):
    """Add tokens to user's balance"""
    try:
        stats = get_stats(user_id)
        if not stats:
            return False

        stats.tokens += amount
        db.session.commit()
        return True
    except Exception as e:
        print(f"Error adding tokens: {e}")
        db.session.rollback()
        return False



def get_subscription_status(user_id):
    """Get user's subscription status"""
    try:
        subscription = Subscription.query.filter_by(
            user_id=user_id,
            active=True
        ).first()

        if subscription:
            # Check if subscription is expired
            if subscription.expires_at and subscription.expires_at < datetime.now(pytz.timezone('Asia/Kolkata')):
                subscription.active = False
                db.session.commit()
                return None

            return {
                'tier': subscription.tier,
                'active': True,
                'expires_at': subscription.expires_at.isoformat() if subscription.expires_at else None
            }

        return None
    except Exception as e:
        print(f"Error getting subscription status: {e}")
        return None

def update_streak_and_xp_fast(user_id, activity_type="login"):
    """Ultra-fast version of streak/XP update optimized for login"""
    try:
        # Get stats directly (skip monthly reset check for speed)
        stats = UserStats.query.filter_by(user_id=user_id).first()
        if not stats:
            print(f"No stats found for user {user_id}")
            return False

        now = datetime.now(pytz.timezone('Asia/Kolkata'))
        today = now.date()

        # Simple streak update (fast path)
        if stats.last_active_date is None:
            stats.streak = 1
        elif stats.last_active_date != today:
            # Only update if it's a new day
            if stats.last_active_date == today - timedelta(days=1):
                stats.streak += 1
            else:
                stats.streak = 1

        stats.last_active_date = today

        # Simple XP update (fast path)
        xp_gains = {'login': 5, 'upload': 15, 'flashcard_generation': 20, 'flashcard_review': 5, 'chat': 10, 'file_chat': 10}
        stats.xp += xp_gains.get(activity_type, 5)
        
        # Simple level calculation
        import math
        new_level = max(1, math.floor(stats.xp / 100) + 1)
        if new_level > stats.level:
            stats.level = new_level
            stats.tokens += 25  # Level up reward

        return True

    except Exception as e:
        print(f"Error in fast streak/XP update: {e}")
        try:
            db.session.rollback()
        except:
            pass
        return False

def update_streak_and_xp(user_id, activity_type="login"):
    """Update user's streak and XP based on activity"""
    try:
        print(f"Updating streak and XP for user {user_id}, activity: {activity_type}")
        stats = get_stats(user_id)  # This includes monthly reset check
        if not stats:
            print(f"Failed to get stats for user {user_id}")
            return False

        now = datetime.now(pytz.timezone('Asia/Kolkata'))
        today = now.date()

        # Handle streak calculation properly
        if stats.last_active_date is None:
            # First time user activity
            stats.streak = 1
        elif stats.last_active_date == today:
            # Already active today, don't change streak
            pass
        elif stats.last_active_date == today - timedelta(days=1):
            # Consecutive day - increment streak
            stats.streak += 1
        else:
            # Missed days - reset streak
            stats.streak = 1

        # Always update last active date
        stats.last_active_date = today

        # Calculate XP gain based on activity type
        xp_gains = {
            'login': 5,
            'upload': 15,
            'flashcard_generation': 20,
            'flashcard_review': 5,
            'chat': 10,
            'file_chat': 10
        }

        xp_gain = xp_gains.get(activity_type, 5)  # Default 5 XP

        # Update XP
        old_level = stats.level
        old_xp = stats.xp
        stats.xp += xp_gain

        # Calculate level: level 1 at 0-99 XP, level 2 at 100-199 XP, etc.
        import math
        stats.level = max(1, math.floor(stats.xp / 100) + 1)

        print(f"XP updated: {old_xp} -> {stats.xp}, Level: {old_level} -> {stats.level}")

        # Check for streak milestone rewards
        streak_reward = 0
        if stats.streak == 7:
            streak_reward = 50
        elif stats.streak == 14:
            streak_reward = 100
        elif stats.streak == 21:
            streak_reward = 150
        elif stats.streak == 30:
            streak_reward = 200
        elif stats.streak > 30 and stats.streak % 30 == 0:
            streak_reward = 200

        # Check for level up reward
        level_up_reward = 0
        if stats.level > old_level:
            level_up_reward = 25  # +25 tokens per level up

        # Apply rewards
        total_tokens_added = streak_reward + level_up_reward
        if total_tokens_added > 0:
            stats.tokens += total_tokens_added

            if streak_reward > 0:
                print(f"🎉 Streak milestone reward: +{streak_reward} tokens for {stats.streak} days")
                track_event('streak_milestone_reward', {
                    'user_id': user_id,
                    'streak_days': stats.streak,
                    'tokens_rewarded': streak_reward
                })

            if level_up_reward > 0:
                print(f"🎊 Level up reward: +{level_up_reward} tokens for reaching level {stats.level}")
                track_event('level_up_reward', {
                    'user_id': user_id,
                    'new_level': stats.level,
                    'tokens_rewarded': level_up_reward
                })

        db.session.commit()
        print(f"✅ Successfully updated stats for user {user_id}: Streak={stats.streak}, XP={stats.xp}, Level={stats.level}, Tokens={stats.tokens}")
        return True

    except Exception as e:
        print(f"Error updating streak and XP: {e}")
        import traceback
        print(traceback.format_exc())
        db.session.rollback()
        return False

def create_referral_code(user_id):
    """Generate a unique referral code for a user"""
    return "IMPIFY-" + user_id[:6].upper()

def apply_referral(code, new_user_id):
    """Apply referral bonus when someone signs up with a referral code"""
    try:
        ref = Referral.query.filter_by(code=code).first()
        if not ref:
            return False

        ref.referee_id = new_user_id
        ref.bonus_given = True

        # Give bonus tokens
        add_tokens(ref.referrer_id, 50)  # Referrer gets 50 tokens
        add_tokens(new_user_id, 20)     # New user gets 20 tokens

        db.session.commit()
        return True
    except Exception as e:
        print(f"Error applying referral: {e}")
        db.session.rollback()
        return False

def can_chat(user_id):
    """Check if user can chat based on free tier limits and tokens"""
    try:
        # Check monthly token reset first
        check_monthly_token_reset(user_id)

        stats = get_stats(user_id)
        if not stats:
            return False

        # Check subscription first (premium users bypass limits)
        subscription = get_subscription_status(user_id)
        if subscription and subscription['tier'] == 'premium':
            return True

        # Check daily free chat limit
        today = date.today()
        chats_today = ChatLog.query.filter(
            ChatLog.user_id == user_id,
            ChatLog.timestamp >= today
        ).count()

        settings = get_global_settings()
        free_limit = settings['free_chats_per_day']

        if chats_today < free_limit:
            return True

        # If user has tokens, allow spending
        return spend_tokens(user_id, 3)  # Cost per chat
    except Exception as e:
        print(f"Error checking chat permission: {e}")
        return False

def check_user_quota(user_id):
    """Check if user has exceeded their free tier quota"""
    ist = pytz.timezone('Asia/Kolkata')
    now = datetime.now(ist)
    today = now.date()
    month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)

    # Get global settings
    settings = get_global_settings()

    # Count today's notes
    today_notes = Note.query.filter(
        Note.user_id == user_id,
        Note.created_at >= today
    ).count()

    # Count this month's notes
    month_notes = Note.query.filter(
        Note.user_id == user_id,
        Note.created_at >= month_start
    ).count()

    # Calculate remaining quota
    daily_remaining = settings['free_uploads_per_day'] - today_notes
    monthly_remaining = settings['free_uploads_per_day'] * 4 - month_notes  # Approximate monthly limit

    return {
        'allowed': daily_remaining > 0 and monthly_remaining > 0,
        'daily_used': today_notes,
        'daily_limit': settings['free_uploads_per_day'],
        'daily_remaining': max(0, daily_remaining),
        'monthly_used': month_notes,
        'monthly_limit': settings['free_uploads_per_day'] * 4,  # Approximate monthly limit
        'monthly_remaining': max(0, monthly_remaining)
    }

@app.route('/api/quota/status', methods=['GET'])
@jwt_required()
@track_usage
def get_quota_status():
    """Get user's current quota status"""
    user_id = get_jwt_identity()
    quota = check_user_quota(user_id)
    return jsonify(quota), 200

# ==================== RAG FUNCTIONS ====================

def chunk_and_embed_text(text, document_id, document_name):
    print("⚠️ Skipped embedding — Vector DB disabled")


def retrieve_relevant_chunks(query, document_id, top_k=5):
    """Retrieve relevant chunks for a document. In fallback mode, return basic text chunks."""
    try:
        # Get the note
        note = Note.query.filter_by(id=document_id).first()
        if not note:
            return []

        # Split the content into chunks
        chunks = text_splitter.split_text(note.content)

        # In fallback mode, return first top_k chunks
        # In production, this would use semantic search with embeddings
        return chunks[:top_k] if chunks else []
    except Exception as e:
        print(f"Error retrieving chunks: {e}")
        return []

# ==================== AI LLM ====================

async def generate_notes(text, document_name, note_type="general", document_id=None):
    system_message = "You are an expert study assistant that creates comprehensive notes."

    # Use RAG if document_id provided
    context_chunks = []
    if document_id:
        # Retrieve relevant chunks based on the note type
        query = f"Create {note_type} notes from {document_name}"
        context_chunks = retrieve_relevant_chunks(query, document_id, top_k=5)
        context = "\n\n".join(context_chunks) if context_chunks else text[:4000]  # fallback to full text
    else:
        context = text[:4000]  # limit for direct processing

    if note_type == "question_paper":
        prompt = f"""
You are an expert tutor analyzing this question paper: {document_name}

Content (questions):
{context}

Your task is to:
1. Answer ALL questions in the paper comprehensively
2. Highlight the most important questions (mark them with ⭐)
3. Provide study suggestions and tips for similar questions

Format your response as:
- **Question 1:** [Question text]
  **Answer:** [Detailed answer]
  ⭐ **Important:** [Why this is important]

- **Question 2:** [Question text]
  **Answer:** [Detailed answer]

**Study Suggestions:**
- [Tip 1]
- [Tip 2]
- [General advice]
"""
    else:
        prompt = f"""
Create structured study notes from: {document_name}

Content:
{context}

Include:
✅ Summary
✅ Key Concepts
✅ Important Points
✅ Formulas
✅ Quick Review
"""

    # Check if user is premium for model selection
    subscription = get_subscription_status(document_id.split('-')[0] if document_id else None)  # Extract user_id from document_id
    is_premium = subscription and subscription['tier'] == 'premium'

    # Try OpenAI first, fallback to Ollama
    # Try OpenAI first, fallback to Ollama, then Gemini
    response = None

    if openai_client:
        try:
            llm_client = LLMClient(provider="openai", use_premium=is_premium)
            response = llm_client.chat(system_message, prompt)
        except Exception as e:
            print(f"OpenAI failed: {e}")

    if response is None:
        try:
            llm_client = LLMClient(provider="ollama", use_premium=is_premium)
            response = llm_client.chat(system_message, prompt)
        except Exception as e:
            print(f"Ollama failed: {e}")

    return response or "AI temporarily unavailable. Try again later."

# ==================== FLASHCARD FUNCTIONS ====================

def generate_basic_flashcards(content, note_title, num_cards=8):
    """Generate basic flashcards when AI is unavailable"""
    try:
        # Split content into sentences
        import re
        sentences = re.split(r'[.!?]+', content)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 20]

        flashcards = []
        for i in range(min(num_cards, len(sentences))):
            sentence = sentences[i]
            # Create a simple question from the sentence
            question = f"What is the main point of: {sentence[:100]}{'...' if len(sentence) > 100 else '?'}"
            answer = sentence

            flashcards.append({
                'question': question,
                'answer': answer
            })

        # If we don't have enough sentences, create generic cards
        while len(flashcards) < num_cards:
            flashcards.append({
                'question': f"What is covered in {note_title}?",
                'answer': f"This document covers various topics related to {note_title}. Please review the content for specific details."
            })

        return flashcards[:num_cards]

    except Exception as e:
        print(f"Error generating basic flashcards: {e}")
        return [
            {
                'question': f"What is {note_title} about?",
                'answer': "Please review the document content for details."
            }
        ]

async def generate_flashcards_from_content(content, note_title, note_type="general", num_cards=8):
    """Generate flashcards from note content using AI"""
    system_message = "You are an expert study assistant that creates effective flashcards for learning and memorization."

    prompt = f"""
    Create {num_cards} high-quality flashcards from the following study material: {note_title}

    Content:
    {content}

    Create flashcards in this exact JSON format:
    [
        {{"question": "Question text here", "answer": "Answer text here"}},
        {{"question": "Question text here", "answer": "Answer text here"}}
    ]

    Guidelines:
    - Questions should be clear, specific, and test important concepts
    - Answers should be concise but complete
    - Mix different types: definitions, concepts, examples, comparisons
    - Focus on key facts, formulas, definitions, and important points
    - Make questions challenging but fair
    - Use the same language as the source material

    Note type: {note_type}
    - If question_paper: Focus on key concepts, formulas, and important points from questions/answers
    - If general: Cover main topics, definitions, and important details

    Return ONLY the JSON array, no additional text.
    """

    # Check if user is premium for model selection
    subscription = get_subscription_status(user_id)
    is_premium = subscription and subscription['tier'] == 'premium'

    # Try OpenAI first, fallback to Ollama
    response = None

    if openai_client:
        try:
            llm_client = LLMClient(provider="openai", use_premium=is_premium)
            response = llm_client.chat(system_message, prompt)
        except Exception as e:
            print(f"OpenAI failed: {e}")

    if response is None:
        try:
            llm_client = LLMClient(provider="ollama", use_premium=is_premium)
            response = llm_client.chat(system_message, prompt)
        except Exception as e:
            print(f"Ollama failed: {e}")

    if not response or "unavailable" in response.lower():
        # Fallback: Generate basic flashcards from content
        print("AI unavailable, generating basic flashcards...")
        return generate_basic_flashcards(content, note_title, num_cards)

    try:
        # Clean response and parse JSON
        import json
        # Remove any markdown formatting
        response = response.strip()
        if response.startswith('```json'):
            response = response[7:]
        if response.endswith('```'):
            response = response[:-3]

        flashcards_data = json.loads(response.strip())

        # Validate the structure
        if not isinstance(flashcards_data, list):
            return generate_basic_flashcards(content, note_title, num_cards)

        # Ensure we have the right format
        validated_cards = []
        for card in flashcards_data:
            if isinstance(card, dict) and 'question' in card and 'answer' in card:
                validated_cards.append({
                    'question': str(card['question']).strip(),
                    'answer': str(card['answer']).strip()
                })

        return validated_cards if validated_cards else generate_basic_flashcards(content, note_title, num_cards)

    except Exception as e:
        print(f"Failed to parse flashcards JSON: {e}")
        print(f"Response was: {response}")
        return generate_basic_flashcards(content, note_title, num_cards)

def calculate_next_review(difficulty_score, review_count, correct_ratio):
    """Calculate next review date using spaced repetition algorithm"""
    import pytz
    from datetime import timedelta
    
    ist = pytz.timezone('Asia/Kolkata')
    now = datetime.now(ist)
    
    # Base interval in days based on difficulty and performance
    if correct_ratio >= 0.8:  # Good performance
        if difficulty_score <= 0.3:  # Easy
            interval = 7 * (review_count + 1)  # 7, 14, 21, 28 days
        elif difficulty_score <= 0.6:  # Medium
            interval = 4 * (review_count + 1)  # 4, 8, 12, 16 days
        else:  # Hard
            interval = 2 * (review_count + 1)  # 2, 4, 6, 8 days
    elif correct_ratio >= 0.5:  # Fair performance
        if difficulty_score <= 0.3:
            interval = 3 * (review_count + 1)  # 3, 6, 9, 12 days
        elif difficulty_score <= 0.6:
            interval = 2 * (review_count + 1)  # 2, 4, 6, 8 days
        else:
            interval = 1 * (review_count + 1)  # 1, 2, 3, 4 days
    else:  # Poor performance
        if difficulty_score <= 0.3:
            interval = 1  # 1 day
        elif difficulty_score <= 0.6:
            interval = 1  # 1 day
        else:
            interval = 0.5  # 12 hours
    
    return now + timedelta(days=interval)

def update_flashcard_progress(flashcard, was_correct, time_taken=None):
    """Update flashcard progress based on user's answer"""
    # Update counters
    flashcard.review_count += 1
    if was_correct:
        flashcard.correct_count += 1
    
    # Calculate performance metrics
    correct_ratio = flashcard.correct_count / flashcard.review_count
    
    # Update difficulty score (0.0 = easy, 1.0 = hard)
    if was_correct:
        # Decrease difficulty slightly
        flashcard.difficulty_score = max(0.0, flashcard.difficulty_score - 0.1)
    else:
        # Increase difficulty
        flashcard.difficulty_score = min(1.0, flashcard.difficulty_score + 0.2)
    
    # Update review dates
    flashcard.last_reviewed = datetime.now(pytz.timezone('Asia/Kolkata'))
    flashcard.next_review = calculate_next_review(flashcard.difficulty_score, flashcard.review_count, correct_ratio)

# ==================== HELPER FUNCTIONS ====================

def extract_text_from_pdf(file_bytes):
    """Extract text from PDF using pdfplumber and PyPDF2 as fallback"""
    text = ""
    
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"pdfplumber failed: {e}, trying PyPDF2...")
        
    if not text.strip():
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"PyPDF2 also failed: {e}")
            return None
    
    return text.strip() if text.strip() else None

def extract_text_from_docx(file_bytes):
    """Extract text from DOCX files using python-docx"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return '\n'.join(text).strip() if text else None
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return None

def extract_text_from_doc(file_bytes):
    """Extract text from DOC files (legacy format) - basic implementation"""
    try:
        # This is a simplified approach for DOC files
        # In production, you might want to use python-docx2txt or other libraries
        import subprocess
        import tempfile
        
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
            temp_file.write(file_bytes)
            temp_file_path = temp_file.name
        
        try:
            # Try using antiword if available (Linux)
            result = subprocess.run(['antiword', temp_file_path],
                                  capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                return result.stdout.strip()
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        
        try:
            # Try using catdoc if available
            result = subprocess.run(['catdoc', temp_file_path],
                                  capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                return result.stdout.strip()
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        
        return None
    except Exception as e:
        print(f"DOC extraction failed: {e}")
        return None
    finally:
        # Clean up temp file
        try:
            os.unlink(temp_file_path)
        except:
            pass

def extract_text_from_txt(file_bytes):
    """Extract text from plain text files"""
    try:
        # Try UTF-8 first
        try:
            return file_bytes.decode('utf-8').strip()
        except UnicodeDecodeError:
            # Try other encodings
            try:
                return file_bytes.decode('latin-1').strip()
            except UnicodeDecodeError:
                return file_bytes.decode('utf-8', errors='ignore').strip()
    except Exception as e:
        print(f"TXT extraction failed: {e}")
        return None

def extract_text_from_image(file_bytes):
    """Extract text from image files using OCR"""
    try:
        from PIL import Image
        import pytesseract
        
        # Open the image
        image = Image.open(io.BytesIO(file_bytes))
        
        # Convert to RGB if needed
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Use pytesseract to extract text
        text = pytesseract.image_to_string(image)
        
        return text.strip() if text.strip() else None
    except Exception as e:
        print(f"Image OCR extraction failed: {e}")
        return None

def extract_text_from_file(file_bytes, filename):
    """Extract text from various file types based on extension"""
    extension = filename.lower().split('.')[-1]

    if extension == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif extension == 'docx':
        return extract_text_from_docx(file_bytes)
    elif extension == 'doc':
        return extract_text_from_doc(file_bytes)
    elif extension in ['txt', 'md', 'text']:
        return extract_text_from_txt(file_bytes)
    elif extension in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp']:
        return extract_text_from_image(file_bytes)
    else:
        return None

def generate_pdf_from_text(title, content):
    """Generate PDF from title and content using ReportLab"""
    if not REPORTLAB_AVAILABLE:
        raise Exception("PDF generation not available")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()

    # Build the PDF content
    story = []

    # Title
    title_style = styles['Heading1']
    title_style.spaceAfter = 30
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 12))

    # Content - split into paragraphs
    content_style = styles['Normal']
    content_style.spaceAfter = 12

    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para.strip(), content_style))
            story.append(Spacer(1, 6))

    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_docx_from_text(title, content):
    """Generate DOCX from title and content using python-docx"""
    try:
        from docx import Document
    except ImportError:
        raise Exception("DOCX generation not available - python-docx not installed")

    doc = Document()
    doc.add_heading(title, 0)

    # Split content into paragraphs
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==================== AUTH ROUTES ====================
@app.route('/api/auth/register', methods=['POST'])
@track_usage
def register():
    try:
        print("=== REGISTRATION ATTEMPT DEBUG ===")
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
        name = data.get('name', '')
        referral_code = data.get('referral_code')  # Optional referral code

        print(f"Registration data - Email: {email}, Name: {name}, Referral: {referral_code}")

        if not email or not password:
            print("❌ Missing email or password")
            return jsonify({"error": "Email and password are required"}), 400

        existing_user = User.query.filter_by(email=email).first()
        if existing_user:
            print(f"❌ User already exists: {email}")
            return jsonify({"error": "User already exists"}), 400

        # Debug password hashing
        print(f"🔑 Raw password before hashing: {password}")
        hashed_password = generate_password_hash(password)
        print(f"🔐 Hashed password: {hashed_password}")

        user = User(
            id=str(uuid.uuid4()),
            email=email,
            name=name,
            password=hashed_password,
            created_at=datetime.now(pytz.timezone('Asia/Kolkata'))
        )

        db.session.add(user)
        db.session.commit()

        # Create referral code for new user
        user_referral_code = create_referral_code(user.id)
        referral = Referral(
            id=str(uuid.uuid4()),
            referrer_id=user.id,
            code=user_referral_code,
            created_at=datetime.now(pytz.timezone('Asia/Kolkata'))
        )
        db.session.add(referral)

        # Apply referral bonus if code was provided
        referral_applied = False
        if referral_code:
            referral_applied = apply_referral(referral_code, user.id)
            if referral_applied:
                print(f"✅ Referral bonus applied for code: {referral_code}")

        db.session.commit()

        # Ensure user stats and subscription exist
        from utils_user_account import ensure_user_stats, ensure_user_subscription
        ensure_user_stats(user.id)
        ensure_user_subscription(user.id)

        print(f"✅ User registered successfully: {email}")

        # Check if user is admin
        is_admin = email == 'admin@gmail.com'
        role = "admin" if is_admin else "user"

        access_token = create_access_token(identity=user.id, additional_claims={"role": role})

        # Track registration
        track_event('user_registered', {
            'email': email,
            'referral_used': bool(referral_code),
            'referral_applied': referral_applied
        })

        return jsonify({
            "message": "User registered successfully",
            "token": access_token,
            "user": {
                "id": user.id,
                "email": user.email,
                "name": user.name,
                "role": role,
                "referral_code": user_referral_code
            },
            "referral_applied": referral_applied
        }), 201

    except Exception as e:
        print(f"❌ Register error: {e}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        db.session.rollback()
        return jsonify({"error": "Registration failed"}), 500
        

@app.route('/api/auth/login', methods=['POST'])
def login():
    """Optimized login - minimal DB operations for fast response"""
    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')

        if not email or not password:
            return jsonify({"error": "Email and password are required"}), 400

        # 1. Find user (REQUIRED - can't skip)
        user = User.query.filter_by(email=email).first()

        if not user:
            return jsonify({"error": "Invalid credentials"}), 401

        # 2. Verify password (REQUIRED - can't skip)
        from werkzeug.security import check_password_hash
        if not check_password_hash(user.password, password):
            return jsonify({"error": "Invalid credentials"}), 401

        # 3. Create token (REQUIRED - can't skip)
        is_admin = email == 'admin@gmail.com'
        role = "admin" if is_admin else "user"
        access_token = create_access_token(identity=user.id, additional_claims={"role": role})

        # 4. Return immediately - NO other DB operations here!
        return jsonify({
            "message": "Login successful",
            "token": access_token,
            "user": {
                "id": user.id,
                "email": user.email,
                "name": user.name or '',
                "role": role
            }
        }), 200

    except Exception as e:
        print(f"Login error: {e}")
        import traceback
        print(traceback.format_exc())
        return jsonify({"error": "Login failed"}), 500

@app.route('/api/admin/auth/login', methods=['POST'])
def admin_login():
    """Admin login endpoint - separate from regular user login"""
    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')

        if not email or not password:
            return jsonify({"error": "Email and password are required"}), 400

        # Check if this is an admin email
        if email != 'admin@gmail.com':
            return jsonify({"error": "Unauthorized access"}), 403

        # Find admin user
        user = User.query.filter_by(email=email).first()

        if not user:
            return jsonify({"error": "Admin user not found"}), 404

        # Verify password
        from werkzeug.security import check_password_hash
        if not check_password_hash(user.password, password):
            return jsonify({"error": "Invalid credentials"}), 401

        # Create admin token with admin role
        access_token = create_access_token(identity=user.id, additional_claims={"role": "admin"})

        return jsonify({
            "message": "Admin login successful",
            "token": access_token,
            "user": {
                "id": user.id,
                "email": user.email,
                "name": user.name or 'Administrator',
                "role": "admin"
            }
        }), 200

    except Exception as e:
        print(f"Admin login error: {e}")
        import traceback
        print(traceback.format_exc())
        return jsonify({"error": "Admin login failed"}), 500


# --- ADD THIS NEW ENDPOINT AFTER THE LOGIN ROUTE ---

@app.route('/api/auth/post-login-init', methods=['POST'])
@jwt_required()
def post_login_init():
    """Initialize user data after login - called separately to not block login"""
    try:
        user_id = get_jwt_identity()
        
        # Ensure user stats exist
        try:
            from utils_user_account import ensure_user_stats, ensure_user_subscription
            ensure_user_stats(user_id)
            ensure_user_subscription(user_id)
            db.session.commit()
        except Exception as e:
            print(f"Failed to ensure user data: {e}")
            try:
                db.session.rollback()
            except:
                pass
        
        # Update streak/XP
        try:
            update_streak_and_xp_fast(user_id, "login")
            db.session.commit()
        except Exception as e:
            print(f"Failed to update streak: {e}")
            try:
                db.session.rollback()
            except:
                pass
        
        # Track login event (non-critical)
        try:
            track_event('user_login', {'user_id': user_id})
        except Exception as e:
            print(f"Failed to track login: {e}")
        return jsonify({"message": "Init complete"}), 200

    except Exception as e:
        print(f"Post-login init error: {e}")
        return jsonify({"error": "Init failed"}), 500


@app.route('/api/subscriptions/plans', methods=['GET'])
def get_subscription_plans():
    """Get all available subscription plans"""
    try:
        plans = list(SUBSCRIPTION_PLANS.values())
        return jsonify({"plans": plans}), 200
    except Exception as e:
        print(f"Get subscription plans error: {e}")
        return jsonify({"error": "Failed to get plans"}), 500


@app.route('/api/subscriptions/current', methods=['GET'])
@jwt_required()
def get_current_subscription():
    """Get user's current subscription"""
    try:
        user_id = get_jwt_identity()

        sub = Subscription.query.filter_by(user_id=user_id, active=True).first()

        if not sub:
            # Return free tier as default
            return jsonify({
                "subscription": {
                    "tier": "free",
                    "active": True,
                    "expires_at": None,
                    "plan": SUBSCRIPTION_PLANS['free']
                }
            }), 200

        plan = SUBSCRIPTION_PLANS.get(sub.tier, SUBSCRIPTION_PLANS['free'])

        return jsonify({
            "subscription": {
                "id": sub.id,
                "tier": sub.tier,
                "active": sub.active,
                "expires_at": sub.expires_at.isoformat() if sub.expires_at else None,
                "created_at": sub.created_at.isoformat() if sub.created_at else None,
                "plan": plan
            }
        }), 200

    except Exception as e:
        print(f"Get current subscription error: {e}")
        return jsonify({"error": "Failed to get subscription"}), 500


@app.route('/api/subscriptions/subscribe', methods=['POST'])
@jwt_required()
def subscribe_to_plan():
    """Initiate subscription to a plan"""
    try:
        user_id = get_jwt_identity()
        data = request.get_json()

        plan_id = data.get('plan_id')

        if not plan_id or plan_id not in SUBSCRIPTION_PLANS:
            return jsonify({"error": "Invalid plan selected"}), 400

        plan = SUBSCRIPTION_PLANS[plan_id]

        # For free plan, just update the subscription directly
        if plan_id == 'free':
            sub = Subscription.query.filter_by(user_id=user_id).first()
            if sub:
                sub.tier = 'free'
                sub.active = True
                sub.expires_at = None
            else:
                sub = Subscription(
                    id=str(uuid.uuid4()),
                    user_id=user_id,
                    tier='free',
                    active=True,
                    created_at=datetime.now(pytz.timezone('Asia/Kolkata')),
                    expires_at=None
                )
                db.session.add(sub)

            db.session.commit()
            return jsonify({
                "message": "Subscription updated to Free plan",
                "subscription": {
                    "tier": "free",
                    "active": True
                }
            }), 200

        # For paid plans, create a pending order (integrate with payment gateway)
        order_id = f"order_{uuid.uuid4().hex[:12]}"

        return jsonify({
            "message": "Order created",
            "order": {
                "order_id": order_id,
                "plan_id": plan_id,
                "plan_name": plan['name'],
                "amount": plan['price'],
                "currency": plan['currency']
            },
            "payment_required": True
        }), 200

    except Exception as e:
        print(f"Subscribe to plan error: {e}")
        return jsonify({"error": "Failed to create subscription"}), 500


@app.route('/api/subscriptions/confirm', methods=['POST'])
@jwt_required()
def confirm_subscription():
    """Confirm subscription after payment"""
    try:
        user_id = get_jwt_identity()
        data = request.get_json()

        order_id = data.get('order_id')
        payment_id = data.get('payment_id')
        plan_id = data.get('plan_id')

        if not order_id or not plan_id:
            return jsonify({"error": "Order ID and plan ID required"}), 400

        if plan_id not in SUBSCRIPTION_PLANS:
            return jsonify({"error": "Invalid plan"}), 400

        plan = SUBSCRIPTION_PLANS[plan_id]

        # In production, verify payment with payment gateway here
        # For now, just activate the subscription

        ist = pytz.timezone('Asia/Kolkata')
        now = datetime.now(ist)
        expires_at = now + timedelta(days=30)  # 30 day subscription

        # Deactivate existing subscription
        existing = Subscription.query.filter_by(user_id=user_id, active=True).first()
        if existing:
            existing.active = False

        # Create new subscription
        new_sub = Subscription(
            id=str(uuid.uuid4()),
            user_id=user_id,
            tier=plan_id,
            active=True,
            created_at=now,
            expires_at=expires_at
        )
        db.session.add(new_sub)

        # Update user stats with new monthly token allocation
        stats = get_stats(user_id)
        if stats and plan['tokens_per_month'] > 0:
            stats.monthly_tokens = plan['tokens_per_month']
            stats.tokens += plan['tokens_per_month']  # Bonus tokens on upgrade

        db.session.commit()

        # Track subscription event
        track_event('subscription_activated', {
            'user_id': user_id,
            'plan_id': plan_id,
            'order_id': order_id
        })

        return jsonify({
            "message": f"Successfully subscribed to {plan['name']} plan!",
            "subscription": {
                "tier": plan_id,
                "active": True,
                "expires_at": expires_at.isoformat()
            }
        }), 200

    except Exception as e:
        print(f"Confirm subscription error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to confirm subscription"}), 500


@app.route('/api/subscriptions/cancel', methods=['POST'])
@jwt_required()
def cancel_subscription():
    """Cancel current subscription"""
    try:
        user_id = get_jwt_identity()

        sub = Subscription.query.filter_by(user_id=user_id, active=True).first()

        if not sub or sub.tier == 'free':
            return jsonify({"error": "No active paid subscription to cancel"}), 400

        # Don't immediately deactivate - let it run until expiry
        # Just mark it for non-renewal
        sub.auto_renew = False if hasattr(sub, 'auto_renew') else None

        db.session.commit()

        return jsonify({
            "message": "Subscription will not renew after current period",
            "expires_at": sub.expires_at.isoformat() if sub.expires_at else None
        }), 200

    except Exception as e:
        print(f"Cancel subscription error: {e}")
        return jsonify({"error": "Failed to cancel subscription"}), 500


@app.route('/api/auth/login-simple', methods=['POST'])
def login_simple():
    """Simplified login for debugging"""
    try:
        print("🔧 SIMPLE LOGIN ATTEMPT")
        
        # Get data
        data = request.get_json() or {}
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        
        print(f"Email: {email}, Password length: {len(password)}")
        
        if not email or not password:
            return jsonify({"error": "Email and password required"}), 400
        
        # Find user
        user = User.query.filter_by(email=email).first()
        if not user:
            return jsonify({"error": "User not found"}), 401
        
        # Check password (direct method)
        from werkzeug.security import check_password_hash
        if check_password_hash(user.password, password):
            # Create token without additional claims first
            token = create_access_token(identity=user.id)
            return jsonify({
                "message": "Login successful",
                "token": token,
                "user": {
                    "id": user.id,
                    "email": user.email,
                    "name": user.name or ""
                }
            }), 200
        else:
            return jsonify({"error": "Invalid password"}), 401
            
    except Exception as e:
        print(f"Simple login error: {e}")
        return jsonify({"error": "Login failed"}), 500
        
        
@app.route('/api/debug/users/<email>', methods=['GET'])
def debug_user_check_route(email):
    """Debug endpoint to check user data"""
    try:
        user = User.query.filter_by(email=email).first()
        if user:
            return jsonify({
                'exists': True,
                'email': user.email,
                'id': user.id,
                'name': user.name,
                'password_hash': user.password,
                'created_at': user.created_at.isoformat() if user.created_at else None
            })
        else:
            return jsonify({'exists': False, 'message': f'User {email} not found'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/auth/verify', methods=['GET'])
@jwt_required()
def verify_auth():
    try:
        user_id = get_jwt_identity()
        user = User.query.get(user_id)

        if not user:
            return jsonify({"valid": False}), 404

        return jsonify({
            "valid": True,
            "user": {
                "id": user.id,
                "email": user.email,
                "name": user.name or '',
                "role": "admin" if user.email == "admin@gmail.com" else "user",
                "consent_given": has_user_consent(user.id)
            }
        }), 200
    except Exception as e:
        print("Verify token error:", e)
        return jsonify({"valid": False}), 401

@app.route('/api/auth/refresh', methods=['POST'])
@jwt_required()
@track_usage
def refresh_token():
    try:
        user_id = get_jwt_identity()
        user = User.query.get(user_id)

        if not user:
            return jsonify({"error": "User not found"}), 404

        # Check if user is admin
        is_admin = user.email == 'admin@gmail.com'
        role = "admin" if is_admin else "user"

        # Create new access token
        access_token = create_access_token(identity=user.id, additional_claims={"role": role})

        # Track token refresh
        track_event('token_refreshed', {'user_id': user_id})

        return jsonify({
            "message": "Token refreshed successfully",
            "token": access_token,
            "user": {
                "id": user.id,
                "email": user.email,
                "name": user.name or '',
                "role": role
            }
        }), 200

    except Exception as e:
        print(f"Token refresh error: {e}")
        return jsonify({"error": "Token refresh failed"}), 500

@app.route('/api/auth/forgot-password', methods=['POST'])
@track_usage
def forgot_password():
    """Send password reset email"""
    try:
        data = request.get_json()
        email = data.get('email')

        if not email:
            return jsonify({"error": "Email is required"}), 400

        user = User.query.filter_by(email=email).first()
        if not user:
            # Don't reveal if email exists or not for security
            return jsonify({"message": "If an account with this email exists, a password reset link has been sent."}), 200

        # Generate reset token
        reset_token = str(uuid.uuid4())
        expires_at = datetime.now(pytz.timezone('Asia/Kolkata')) + timedelta(hours=1)  # 1 hour expiry

        # Save reset token
        password_reset = PasswordResetToken(
            id=str(uuid.uuid4()),
            user_id=user.id,
            token=reset_token,
            expires_at=expires_at
        )

        db.session.add(password_reset)
        db.session.commit()

        # In a real application, you would send an email here
        # For now, we'll just return the token in the response for testing
        reset_link = f"https://impify.visasystem.in/reset-password?token={reset_token}"

        print(f"Password reset link for {email}: {reset_link}")

        return jsonify({
            "message": "Password reset link sent to your email.",
            "reset_link": reset_link  # Remove this in production
        }), 200

    except Exception as e:
        print(f"Forgot password error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to process request"}), 500

@app.route('/api/auth/reset-password', methods=['POST'])
@track_usage
def reset_password():
    """Reset password using token"""
    try:
        data = request.get_json()
        token = data.get('token')
        new_password = data.get('new_password')

        if not token or not new_password:
            return jsonify({"error": "Token and new password are required"}), 400

        if len(new_password) < 6:
            return jsonify({"error": "Password must be at least 6 characters long"}), 400

        # Find valid reset token
        reset_token = PasswordResetToken.query.filter_by(
            token=token,
            used=False
        ).first()

        if not reset_token:
            return jsonify({"error": "Invalid or expired reset token"}), 400

        # Check if token is expired
        if reset_token.expires_at < datetime.now(pytz.timezone('Asia/Kolkata')):
            return jsonify({"error": "Reset token has expired"}), 400

        # Update user password
        user = reset_token.user
        user.password = generate_password_hash(new_password)

        # Mark token as used
        reset_token.used = True

        db.session.commit()

        return jsonify({"message": "Password reset successfully"}), 200

    except Exception as e:
        print(f"Reset password error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to reset password"}), 500

@app.route('/api/referral/validate', methods=['POST'])
def validate_referral_code():
    """Validate if a referral code exists and is available"""
    try:
        data = request.get_json()
        code = data.get('code')

        if not code:
            return jsonify({"error": "Referral code is required"}), 400

        # Check if referral code exists and hasn't been used
        referral = Referral.query.filter_by(code=code).first()

        if not referral:
            return jsonify({
                "valid": False,
                "message": "Invalid referral code"
            }), 200

        # Check if code has already been used
        if referral.referee_id:
            return jsonify({
                "valid": False,
                "message": "Referral code has already been used"
            }), 200

        return jsonify({
            "valid": True,
            "message": "Valid referral code"
        }), 200

    except Exception as e:
        print(f"Validate referral code error: {e}")
        return jsonify({"error": "Failed to validate referral code"}), 500



        
        

@app.route('/api/debug/test-password', methods=['POST'])
def test_password_hashing():
    """Test if password hashing works correctly"""
    try:
        data = request.get_json()
        password = data.get('password')
        
        if not password:
            return jsonify({'error': 'Password required'}), 400
            
        # Test hashing
        hashed = generate_password_hash(password)
        check_result = check_password_hash(hashed, password)
        
        return jsonify({
            'input_password': password,
            'hashed_password': hashed,
            'verification_result': check_result
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500
        
        
        
# ==================== DEBUG ROUTES ====================

@app.route('/api/debug/test-hashing', methods=['GET'])
def test_hashing():
    """Test password hashing functionality"""
    from werkzeug.security import generate_password_hash, check_password_hash
    
    test_password = "test123"
    hashed = generate_password_hash(test_password)
    check_result = check_password_hash(hashed, test_password)
    
    return jsonify({
        'werkzeug_version': '3.1.3',
        'flask_version': '2.3.3',
        'test_password': test_password,
        'hashed': hashed,
        'verification_result': check_result,
        'success': check_result
    })

       
@app.route('/api/debug/emergency-fix', methods=['POST'])
def emergency_fix():
    """Emergency fix for existing users"""
    try:
        data = request.get_json()
        email = data.get('email')
        new_password = data.get('new_password')
        
        user = User.query.filter_by(email=email).first()
        if user:
            # Use the proper hashing method
            from werkzeug.security import generate_password_hash
            user.password = generate_password_hash(new_password)
            db.session.commit()
            
            # Verify the fix
            from werkzeug.security import check_password_hash
            verify = check_password_hash(user.password, new_password)
            
            return jsonify({
                'message': f'Password updated for {email}',
                'verification_result': verify,
                'success': verify
            })
        else:
            return jsonify({'error': 'User not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ==================== NOTES ROUTES ====================

@app.route('/api/notes/upload', methods=['POST'])
@jwt_required()
@track_usage
def upload_document():
    start_time = time.time()
    try:
        user_id = get_jwt_identity()
        
        # Check quota FIRST before processing
        quota = check_user_quota(user_id)
        if not quota['allowed']:
            track_event('upload_blocked', {
                'reason': 'quota_exceeded',
                'daily_used': quota['daily_used'],
                'monthly_used': quota['monthly_used']
            })
            
            return jsonify({
                "error": "Free tier limit reached",
                "message": f"You've reached your free tier limit. Daily: {quota['daily_used']}/{quota['daily_limit']}, Monthly: {quota['monthly_used']}/{quota['monthly_limit']}",
                "quota": quota
            }), 429  # Too Many Requests
        
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        note_type = request.form.get('note_type', 'general')
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # Support multiple file types
        allowed_extensions = ['.pdf', '.docx', '.doc', '.txt', '.md', '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']
        file_extension = '.' + file.filename.lower().split('.')[-1]
        
        if file_extension not in allowed_extensions:
            return jsonify({
                "error": "Unsupported file type",
                "supported_types": "PDF, DOC, DOCX, TXT, MD, JPG, PNG, GIF, BMP, TIFF, WEBP"
            }), 400
        
        # Track upload start
        track_event('file_upload_started', {
            'filename': file.filename,
            'note_type': note_type
        })
        
        file_bytes = file.read()
        file_size = len(file_bytes)
        file_size_mb = file_size / (1024 * 1024)

        # Check file size limit (using default for now - in production this would be configurable)
        max_file_size_mb = 10  # Default max file size
        if file_size_mb > max_file_size_mb:
            return jsonify({
                "error": f"File too large. Maximum file size is {max_file_size_mb}MB. Your file is {file_size_mb:.1f}MB"
            }), 413  # Payload Too Large
        
        # Extract text using the new multi-format function
        extracted_text = extract_text_from_file(file_bytes, file.filename)

        if not extracted_text:
            track_event('file_upload_failed', {
                'filename': file.filename,
                'reason': 'text_extraction_failed'
            })
            return jsonify({"error": f"Could not extract text from {file_extension.upper()} file."}), 400

        # Generate unique document ID
        document_id = str(uuid.uuid4())

        # Chunk and embed the text for RAG
        try:
            chunk_and_embed_text(extracted_text, document_id, file.filename)
        except Exception as e:
            print(f"Embedding failed: {e}")
            # Continue without RAG if embedding fails

        # Generate notes using AI with RAG
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        generated_notes = loop.run_until_complete(
            generate_notes(extracted_text, file.filename, note_type, document_id)
        )
        loop.close()

        if not generated_notes or "unavailable" in generated_notes.lower():
            track_event('ai_generation', {'success': False, 'reason': 'ai_unavailable'})
            return jsonify({"error": "AI could not process this file right now. Please try again later."}), 500

        # Track successful AI generation
        track_event('ai_generation', {'success': True, 'note_type': note_type})

        # Save note
        note = Note(
            id=document_id,  # Use same ID as document for linking
            user_id=user_id,
            title=file.filename.rsplit('.', 1)[0],  # Remove extension from any file type
            original_filename=file.filename,
            note_type=note_type,
            content=generated_notes,
            file_size=file_size,
            processing_time=round(time.time() - start_time, 2),
            created_at=datetime.now(pytz.timezone('Asia/Kolkata')),
            updated_at=datetime.now(pytz.timezone('Asia/Kolkata'))
        )

        db.session.add(note)
        db.session.commit()

        # Log training example
        log_training_example(
            user_id=user_id,
            source_file=file.filename,
            prompt=f"Create {note_type} notes from {file.filename}",
            input_text=extracted_text[:4000],
            output_text=generated_notes,
            meta={
                "note_type": note_type,
                "file_size": file_size,
                "processing_time": note.processing_time,
                "chunks_embedded": len(text_splitter.split_text(extracted_text))
            }
        )

        # Update user streak and XP for file upload activity
        update_streak_and_xp(user_id, "upload")

        # Track successful upload and generation
        track_event('note_generated', {
            'note_id': note.id,
            'note_type': note_type,
            'file_size': file_size,
            'processing_time': note.processing_time,
            'chunks_embedded': len(text_splitter.split_text(extracted_text))
        })

        return jsonify({
            "message": "Notes generated successfully",
            "note": {
                "id": note.id,
                "title": note.title,
                "note_type": note.note_type,
                "content": note.content,
                "created_at": note.created_at.isoformat()
            }
        }), 201
    except Exception as e:
        print(f"Upload error: {e}")
        track_event('note_generation_failed', {'error': str(e)})
        return jsonify({"error": f"Upload failed: {str(e)}"}), 500

@app.route('/api/notes', methods=['GET'])
@jwt_required()
@track_usage
def get_notes():
    try:
        user_id = get_jwt_identity()
        notes = Note.query.filter_by(user_id=user_id).order_by(Note.created_at.desc()).all()

        notes_data = []
        for note in notes:
            notes_data.append({
                "id": note.id,
                "title": note.title,
                "original_filename": note.original_filename,
                "note_type": note.note_type,
                "content": note.content,
                "file_size": note.file_size,
                "processing_time": note.processing_time,
                "created_at": note.created_at.isoformat(),
                "updated_at": note.updated_at.isoformat()
            })

        track_event('notes_viewed', {'count': len(notes_data)})
        return jsonify({"notes": notes_data}), 200
    except Exception as e:
        print(f"Get notes error: {e}")
        return jsonify({"error": "Failed to fetch notes"}), 500

@app.route('/api/notes/<note_id>', methods=['GET'])
@jwt_required()
@track_usage
def get_note(note_id):
    try:
        user_id = get_jwt_identity()
        note = Note.query.filter_by(id=note_id, user_id=user_id).first()

        if not note:
            return jsonify({"error": "Note not found"}), 404

        note_data = {
            "id": note.id,
            "title": note.title,
            "original_filename": note.original_filename,
            "note_type": note.note_type,
            "content": note.content,
            "file_size": note.file_size,
            "processing_time": note.processing_time,
            "created_at": note.created_at.isoformat(),
            "updated_at": note.updated_at.isoformat()
        }

        track_event('note_opened', {'note_id': note_id})
        return jsonify({"note": note_data}), 200
    except Exception as e:
        print(f"Get note error: {e}")
        return jsonify({"error": "Failed to fetch note"}), 500

@app.route('/api/notes/<note_id>', methods=['DELETE'])
@jwt_required()
@track_usage
def delete_note(note_id):
    try:
        user_id = get_jwt_identity()
        note = Note.query.filter_by(id=note_id, user_id=user_id).first()

        if not note:
            return jsonify({"error": "Note not found"}), 404

        db.session.delete(note)
        db.session.commit()

        track_event('note_deleted', {'note_id': note_id})
        return jsonify({"message": "Note deleted successfully"}), 200
    except Exception as e:
        print(f"Delete note error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to delete note"}), 500

@app.route('/api/notes/<note_id>/export', methods=['GET'])
@jwt_required()
@track_usage
def export_note(note_id):
    try:
        user_id = get_jwt_identity()
        export_format = request.args.get('format', 'txt')

        note = Note.query.filter_by(id=note_id, user_id=user_id).first()

        if not note:
            return jsonify({"error": "Note not found"}), 404

        if export_format == 'pdf':
            # PDF export disabled on shared hosting
            return jsonify({"error": "PDF export is not available on this server. Please use TXT export instead."}), 400
        elif export_format == 'docx':
            try:
                output = generate_docx_from_text(note.title, note.content)
                mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                filename = f"{note.title}.docx"
            except Exception as e:
                return jsonify({"error": f"DOCX export failed: {str(e)}"}), 500
        else:
            # Default to TXT
            content = f"Title: {note.title}\n\n{note.content}"
            output = io.BytesIO()
            output.write(content.encode('utf-8'))
            output.seek(0)
            mimetype = 'text/plain'
            filename = f"{note.title}.txt"

        track_event('note_exported', {
            'note_id': note_id,
            'format': export_format
        })

        return send_file(
            output,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        print(f"Export error: {e}")
        return jsonify({"error": "Failed to export note"}), 500

# ==================== FLASHCARD ROUTES ====================

@app.route('/api/flashcards/generate', methods=['POST'])
@jwt_required()
@track_usage
def generate_flashcards():
    """Generate flashcards from a note using AI"""
    try:
        user_id = get_jwt_identity()
        if not user_id:
            return jsonify({"error": "User not authenticated"}), 401
        
        data = request.get_json()
        if not data:
            return jsonify({"error": "Request data is required"}), 400
            
        note_id = data.get('note_id')
        num_cards = data.get('num_cards', 8)
        
        if not note_id:
            return jsonify({"error": "note_id is required"}), 400
        
        # Get the note
        note = Note.query.filter_by(id=note_id, user_id=user_id).first()
        if not note:
            return jsonify({"error": "Note not found"}), 404
        
        # Check if flashcards already exist for this note
        existing_cards = Flashcard.query.filter_by(note_id=note_id, user_id=user_id).count()
        if existing_cards > 0:
            return jsonify({"error": "Flashcards already exist for this note"}), 400
        
        # Check if note has content
        if not note.content or not note.content.strip():
            return jsonify({"error": "Note has no content to generate flashcards from"}), 400
        
        # Generate flashcards using AI
        try:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            flashcards_data = loop.run_until_complete(
                generate_flashcards_from_content(note.content, note.title, note.note_type, num_cards)
            )
            loop.close()
        except Exception as ai_error:
            print(f"AI generation error: {ai_error}")
            return jsonify({"error": "AI service temporarily unavailable"}), 503
        
        if not flashcards_data:
            return jsonify({"error": "Failed to generate flashcards - AI returned empty result"}), 500
        
        # Create flashcard records
        created_cards = []
        try:
            for card_data in flashcards_data:
                if not card_data or 'question' not in card_data or 'answer' not in card_data:
                    print(f"Skipping invalid card data: {card_data}")
                    continue
                    
                flashcard = Flashcard(
                    id=str(uuid.uuid4()),
                    user_id=user_id,
                    note_id=note_id,
                    question=str(card_data['question']).strip(),
                    answer=str(card_data['answer']).strip(),
                    created_at=datetime.now(pytz.timezone('Asia/Kolkata')),
                    updated_at=datetime.now(pytz.timezone('Asia/Kolkata'))
                )
                db.session.add(flashcard)
                created_cards.append(flashcard)
            
            db.session.commit()
        except Exception as db_error:
            print(f"Database error: {db_error}")
            db.session.rollback()
            return jsonify({"error": "Failed to save flashcards to database"}), 500

        # Update user streak and XP for flashcard generation activity
        update_streak_and_xp(user_id, "flashcard_generation")

        # Track flashcard generation
        try:
            track_event('flashcards_generated', {
                'note_id': note_id,
                'count': len(created_cards)
            })
        except Exception as track_error:
            print(f"Tracking error: {track_error}")
            # Don't fail the request if tracking fails

        return jsonify({
            "message": f"Generated {len(created_cards)} flashcards successfully",
            "flashcards": [{
                "id": card.id,
                "question": card.question,
                "answer": card.answer
            } for card in created_cards]
        }), 201
        
    except Exception as e:
        print(f"Generate flashcards error: {e}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        db.session.rollback()
        return jsonify({"error": f"Failed to generate flashcards: {str(e)}"}), 500

@app.route('/api/flashcards', methods=['GET'])
@jwt_required()
@track_usage
def get_flashcards():
    """Get user's flashcards, optionally filtered by note_id"""
    try:
        user_id = get_jwt_identity()
        note_id = request.args.get('note_id')
        
        if not user_id:
            return jsonify({"error": "User not authenticated"}), 401
        
        query = Flashcard.query.filter_by(user_id=user_id)
        
        if note_id:
            query = query.filter_by(note_id=note_id)
        
        flashcards = query.order_by(Flashcard.created_at.desc()).all()
        
        flashcards_data = []
        for card in flashcards:
            # Safely handle potential None values
            try:
                # Get note title if available
                note_title = None
                if card.note_id:
                    note = Note.query.filter_by(id=card.note_id).first()
                    note_title = note.title if note else None

                flashcards_data.append({
                    "id": card.id,
                    "note_id": card.note_id,
                    "note_title": note_title,
                    "question": card.question or "",
                    "answer": card.answer or "",
                    "difficulty_score": card.difficulty_score or 0.0,
                    "review_count": card.review_count or 0,
                    "correct_count": card.correct_count or 0,
                    "correct_ratio": round((card.correct_count or 0) / max(card.review_count or 1, 1), 2),
                    "next_review": card.next_review.isoformat() if card.next_review else None,
                    "created_at": card.created_at.isoformat() if card.created_at else None
                })
            except Exception as card_error:
                print(f"Error processing flashcard {card.id}: {card_error}")
                # Skip problematic flashcards
                continue
        
        return jsonify({"flashcards": flashcards_data}), 200
        
    except Exception as e:
        print(f"Get flashcards error: {e}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return jsonify({"error": f"Failed to fetch flashcards: {str(e)}"}), 500

@app.route('/api/flashcards/due', methods=['GET'])
@jwt_required()
@track_usage
def get_due_flashcards():
    """Get flashcards due for review using spaced repetition"""
    try:
        user_id = get_jwt_identity()
        if not user_id:
            return jsonify({"error": "User not authenticated"}), 401
        
        limit = request.args.get('limit', 20)  # Default to 20 cards
        try:
            limit = int(limit)
        except ValueError:
            limit = 20
        
        now = datetime.now(pytz.timezone('Asia/Kolkata'))
        
        # Get cards due for review (next_review <= now) and cards that have never been reviewed
        due_cards = Flashcard.query.filter(
            Flashcard.user_id == user_id,
            (Flashcard.next_review <= now) | (Flashcard.next_review.is_(None))
        ).order_by(
            # Order by difficulty (harder cards first) and then by next_review
            Flashcard.difficulty_score.desc(),
            Flashcard.next_review.asc()
        ).limit(limit).all()
        
        flashcards_data = []
        for card in due_cards:
            try:
                flashcards_data.append({
                    "id": card.id,
                    "note_id": card.note_id,
                    "question": card.question or "",
                    "answer": card.answer or "",
                    "difficulty_score": card.difficulty_score or 0.0,
                    "review_count": card.review_count or 0,
                    "correct_count": card.correct_count or 0,
                    "correct_ratio": round((card.correct_count or 0) / max(card.review_count or 1, 1), 2),
                    "is_new": (card.review_count or 0) == 0
                })
            except Exception as card_error:
                print(f"Error processing flashcard {card.id}: {card_error}")
                continue
        
        return jsonify({"flashcards": flashcards_data}), 200
        
    except Exception as e:
        print(f"Get due flashcards error: {e}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return jsonify({"error": f"Failed to fetch due flashcards: {str(e)}"}), 500

@app.route('/api/flashcards/<flashcard_id>/review', methods=['PATCH'])
@jwt_required()
@track_usage
def review_flashcard(flashcard_id):
    """Update flashcard progress after user review"""
    try:
        user_id = get_jwt_identity()
        data = request.get_json()
        was_correct = data.get('was_correct', False)
        
        flashcard = Flashcard.query.filter_by(id=flashcard_id, user_id=user_id).first()
        if not flashcard:
            return jsonify({"error": "Flashcard not found"}), 404
        
        # Update progress
        update_flashcard_progress(flashcard, was_correct)
        db.session.commit()

        # Update user streak and XP for flashcard review activity
        update_streak_and_xp(user_id, "flashcard_review")

        # Track review
        track_event('flashcard_reviewed', {
            'flashcard_id': flashcard_id,
            'was_correct': was_correct,
            'difficulty_score': flashcard.difficulty_score
        })

        return jsonify({
            "message": "Flashcard updated successfully",
            "flashcard": {
                "id": flashcard.id,
                "difficulty_score": flashcard.difficulty_score,
                "review_count": flashcard.review_count,
                "correct_count": flashcard.correct_count,
                "correct_ratio": round(flashcard.correct_count / flashcard.review_count, 2) if flashcard.review_count > 0 else 0,
                "next_review": flashcard.next_review.isoformat() if flashcard.next_review else None
            }
        }), 200
        
    except Exception as e:
        print(f"Review flashcard error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to update flashcard"}), 500

@app.route('/api/flashcards/<flashcard_id>', methods=['DELETE'])
@jwt_required()
@track_usage
def delete_flashcard(flashcard_id):
    """Delete a flashcard"""
    try:
        user_id = get_jwt_identity()
        flashcard = Flashcard.query.filter_by(id=flashcard_id, user_id=user_id).first()
        
        if not flashcard:
            return jsonify({"error": "Flashcard not found"}), 404
        
        db.session.delete(flashcard)
        db.session.commit()
        
        track_event('flashcard_deleted', {'flashcard_id': flashcard_id})
        return jsonify({"message": "Flashcard deleted successfully"}), 200
        
    except Exception as e:
        print(f"Delete flashcard error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to delete flashcard"}), 500


# ==================== COMMUNITY ROUTES ====================

@app.route('/api/community/posts', methods=['GET'])
@jwt_required()
@track_usage
def get_community_posts():
    try:
        user_id = get_jwt_identity()

        # Get all community posts from regular users
        posts = CommunityPost.query.order_by(CommunityPost.created_at.desc()).limit(50).all()

        posts_data = []
        for post in posts:
            # Check if current user liked this post
            is_liked = CommunityLike.query.filter_by(
                post_id=post.id,
                user_id=user_id
            ).first() is not None

            likes_count = CommunityLike.query.filter_by(post_id=post.id).count()

            posts_data.append({
                "id": post.id,
                "content": post.content,
                "created_at": post.created_at.isoformat(),
                "author": {
                    "id": post.author_id,
                    "name": post.author_name,
                    "email": post.author_email,
                    "role": "Community Member"
                },
                "likes_count": likes_count,
                "is_liked": is_liked
            })

        return jsonify({"posts": posts_data}), 200
    except Exception as e:
        print(f"Get community posts error: {e}")
        return jsonify({"error": "Failed to fetch community posts"}), 500

@app.route('/api/community/posts', methods=['POST'])
@jwt_required()
@track_usage
def create_community_post():
    try:
        user_id = get_jwt_identity()
        data = request.get_json()

        if not data.get('content'):
            return jsonify({"error": "Content is required"}), 400

        # Get user details from database
        user = User.query.filter_by(id=user_id).first()
        if not user:
            return jsonify({"error": "User not found"}), 404

        post = CommunityPost(
            id=str(uuid.uuid4()),
            author_id=user_id,
            author_name=user.name or 'Anonymous',
            author_email=user.email,
            content=data['content'],
            created_at=datetime.now(pytz.timezone('Asia/Kolkata')),
            updated_at=datetime.now(pytz.timezone('Asia/Kolkata'))
        )

        db.session.add(post)
        db.session.commit()

        # Track post creation
        track_event('community_post_created', {
            'post_id': post.id,
            'author_id': user_id
        })

        return jsonify({
            "message": "Post created successfully",
            "post": {
                "id": post.id,
                "content": post.content,
                "created_at": post.created_at.isoformat(),
                "author": {
                    "id": post.author_id,
                    "name": post.author_name,
                    "email": post.author_email,
                    "role": "Community Member"
                },
                "likes_count": 0,
                "is_liked": False
            }
        }), 201
    except Exception as e:
        print(f"Create community post error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to create post"}), 500

@app.route('/api/community/posts/<post_id>/like', methods=['POST'])
@jwt_required()
@track_usage
def like_community_post(post_id):
    try:
        user_id = get_jwt_identity()

        # Check if post exists
        post = CommunityPost.query.filter_by(id=post_id).first()
        if not post:
            return jsonify({"error": "Post not found"}), 404

        # Check if already liked
        existing_like = CommunityLike.query.filter_by(
            post_id=post_id,
            user_id=user_id
        ).first()

        if existing_like:
            # Unlike
            db.session.delete(existing_like)
            action = 'unliked'
        else:
            # Like
            like = CommunityLike(
                id=str(uuid.uuid4()),
                post_id=post_id,
                user_id=user_id,
                created_at=datetime.now(pytz.timezone('Asia/Kolkata'))
            )
            db.session.add(like)
            action = 'liked'

        db.session.commit()

        # Get updated likes count
        likes_count = CommunityLike.query.filter_by(post_id=post_id).count()

        # Track like/unlike
        track_event('community_post_liked', {
            'post_id': post_id,
            'user_id': user_id,
            'action': action
        })

        return jsonify({
            "message": f"Post {action}",
            "likes_count": likes_count
        }), 200
    except Exception as e:
        print(f"Like community post error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to like post"}), 500

@app.route('/api/community/stats', methods=['GET'])
@jwt_required()
@track_usage
def get_community_stats():
    try:
        # Get total registered users (from User model)
        total_members = User.query.count()

        # Get total posts
        total_posts = CommunityPost.query.count()

        # Get total likes
        total_likes = CommunityLike.query.count()

        # Get top contributors (users with most posts)
        from sqlalchemy import func
        top_contributors_query = db.session.query(
            CommunityPost.author_name,
            func.count(CommunityPost.id).label('post_count')
        ).group_by(CommunityPost.author_name).order_by(
            func.count(CommunityPost.id).desc()
        ).limit(5).all()

        top_contributors = len([c for c in top_contributors_query if c.post_count > 0])

        return jsonify({
            "total_members": total_members,
            "total_posts": total_posts,
            "total_likes": total_likes,
            "top_contributors": top_contributors
        }), 200
    except Exception as e:
        print(f"Get community stats error: {e}")
        return jsonify({"error": "Failed to fetch community stats"}), 500

@app.route('/api/support/tickets', methods=['GET'])
@jwt_required()
@track_usage
def get_support_tickets():
    try:
        user_id = get_jwt_identity()

        # Get tickets created by the current user
        tickets = SupportTicket.query.filter_by(author_id=user_id).order_by(SupportTicket.created_at.desc()).all()

        tickets_data = []
        for ticket in tickets:
            tickets_data.append({
                "id": ticket.id,
                "title": ticket.title,
                "description": ticket.description,
                "category": ticket.category,
                "priority": ticket.priority,
                "status": ticket.status,
                "author_id": ticket.author_id,
                "created_at": ticket.created_at.isoformat(),
                "updated_at": ticket.updated_at.isoformat()
            })

        return jsonify({"tickets": tickets_data}), 200
    except Exception as e:
        print(f"Get support tickets error: {e}")
        return jsonify({"error": "Failed to fetch support tickets"}), 500

@app.route('/api/support/tickets', methods=['POST'])
@jwt_required()
@track_usage
def create_support_ticket():
    try:
        user_id = get_jwt_identity()
        data = request.get_json()

        # Validate required fields
        required_fields = ['title', 'description', 'category']
        for field in required_fields:
            if not data.get(field):
                return jsonify({"error": f"{field} is required"}), 400

        # Validate category and priority
        valid_categories = ['bug', 'feature', 'improvement', 'question', 'other']
        valid_priorities = ['low', 'medium', 'high', 'urgent']

        if data['category'] not in valid_categories:
            return jsonify({"error": "Invalid category"}), 400

        priority = data.get('priority', 'medium')
        if priority not in valid_priorities:
            return jsonify({"error": "Invalid priority"}), 400

        ticket = SupportTicket(
            id=str(uuid.uuid4()),
            title=data['title'],
            description=data['description'],
            category=data['category'],
            priority=priority,
            status='open',
            author_id=user_id,
            created_at=datetime.now(pytz.timezone('Asia/Kolkata')),
            updated_at=datetime.now(pytz.timezone('Asia/Kolkata'))
        )

        db.session.add(ticket)
        db.session.commit()

        # Track ticket creation
        track_event('support_ticket_created', {
            'ticket_id': ticket.id,
            'category': ticket.category,
            'priority': ticket.priority
        })

        return jsonify({
            "message": "Support ticket created successfully",
            "ticket": {
                "id": ticket.id,
                "title": ticket.title,
                "description": ticket.description,
                "category": ticket.category,
                "priority": ticket.priority,
                "status": ticket.status,
                "author_id": ticket.author_id,
                "created_at": ticket.created_at.isoformat(),
                "updated_at": ticket.updated_at.isoformat()
            }
        }), 201
    except Exception as e:
        print(f"Create support ticket error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to create support ticket"}), 500


@app.route('/api/support/stats', methods=['GET'])
@jwt_required()
@track_usage
def get_support_stats():
    try:
        user_id = get_jwt_identity()

        # Get user's ticket statistics
        total_tickets = SupportTicket.query.filter_by(author_id=user_id).count()
        open_tickets = SupportTicket.query.filter_by(author_id=user_id, status='open').count()
        in_progress_tickets = SupportTicket.query.filter_by(author_id=user_id, status='in_progress').count()
        resolved_tickets = SupportTicket.query.filter_by(author_id=user_id, status='resolved').count()

        return jsonify({
            "total_tickets": total_tickets,
            "open_tickets": open_tickets,
            "in_progress_tickets": in_progress_tickets,
            "resolved_tickets": resolved_tickets
        }), 200
    except Exception as e:
        print(f"Get support stats error: {e}")
        return jsonify({"error": "Failed to fetch support stats"}), 500

# ==================== ADMIN AUTH ROUTES ====================
@app.route('/api/admin/auth/login', methods=['POST'])
def admin_login():
    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')

        if not email or not password:
            return jsonify({"error": "Email and password are required"}), 400

        user = User.query.filter_by(email=email).first()

        if not user or not check_password_hash(user.password, password):
            track_event('login_failed', {'email': email})
            return jsonify({"error": "Invalid credentials"}), 401

        # Must have admin email
        if email != 'admin@gmail.com':
            return jsonify({"error": "Admin access denied"}), 403

        # FIX ✅ Ensure role is included
        role = "admin"
        access_token = create_access_token(identity=user.id, additional_claims={"role": role})

        track_event('admin_login', {'email': email})

        return jsonify({
            "message": "Admin login successful",
            "token": access_token,
            "user": {
                "id": user.id,
                "email": user.email,
                "name": user.name or '',
                "role": role  # ✅ FIX HERE
            }
        }), 200

    except Exception as e:
        print(f"Admin login error: {e}")
        return jsonify({"error": "Login failed"}), 500


# ==================== ADMIN DASHBOARD ROUTES ====================

@app.route('/api/admin/dashboard/stats', methods=['GET'])
@jwt_required()
@admin_required
def get_admin_dashboard_stats():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        # Get total users
        total_users = User.query.count()

        # Get active users (users who have logged in within last 30 days)
        thirty_days_ago = datetime.now(pytz.timezone('Asia/Kolkata')) - timedelta(days=30)
        active_users = Analytics.query.filter(
            Analytics.event_type == 'user_login',
            Analytics.timestamp >= thirty_days_ago
        ).distinct(Analytics.user_id).count()

        # Get total notes
        total_notes = Note.query.count()

        # Get total uploads
        total_uploads = Analytics.query.filter_by(event_type="file_upload_started").count()

        # Get pending volunteer applications (for now, count all users as potential volunteers)
        pending_volunteers = User.query.count()  # In a real app, you'd have a separate volunteer application table

        # Get recent activity (last 7 days)
        seven_days_ago = datetime.now(pytz.timezone('Asia/Kolkata')) - timedelta(days=7)
        recent_uploads = Analytics.query.filter(
            Analytics.event_type == 'file_upload_started',
            Analytics.timestamp >= seven_days_ago
        ).count()

        recent_registrations = User.query.filter(
            User.created_at >= seven_days_ago
        ).count()

        return jsonify({
            "total_users": total_users,
            "active_users": active_users,
            "total_notes": total_notes,
            "total_uploads": total_uploads,
            "pending_volunteers": pending_volunteers,
            "recent_activity": {
                "uploads": recent_uploads,
                "registrations": recent_registrations
            }
        }), 200
    except Exception as e:
        print(f"Get admin dashboard stats error: {e}")
        return jsonify({"error": "Failed to fetch dashboard stats"}), 500

@app.route('/api/admin/dashboard/recent-activity', methods=['GET'])
@jwt_required()
@admin_required
def get_admin_recent_activity():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        # Get recent uploads
        recent_uploads = Analytics.query.filter_by(event_type="file_upload_started").order_by(
            Analytics.timestamp.desc()
        ).limit(10).all()

        # Get recent registrations
        recent_registrations = User.query.order_by(User.created_at.desc()).limit(10).all()

        activity_data = []

        # Add uploads to activity
        for upload in recent_uploads:
            user = User.query.filter_by(id=upload.user_id).first()
            activity_data.append({
                "id": upload.event_id,
                "type": "upload",
                "description": f"File uploaded by {user.name if user else 'Unknown user'}",
                "timestamp": upload.timestamp.isoformat(),
                "user": user.name if user else 'Unknown'
            })

        # Add registrations to activity
        for registration in recent_registrations:
            activity_data.append({
                "id": registration.id,
                "type": "registration",
                "description": f"New user registered: {registration.name}",
                "timestamp": registration.created_at.isoformat(),
                "user": registration.name
            })

        # Sort by timestamp and limit to 20 most recent
        activity_data.sort(key=lambda x: x['timestamp'], reverse=True)
        activity_data = activity_data[:20]

        return jsonify({"activities": activity_data}), 200
    except Exception as e:
        print(f"Get admin recent activity error: {e}")
        return jsonify({"error": "Failed to fetch recent activity"}), 500

# ==================== ADMIN USER MANAGEMENT ROUTES ====================

@app.route('/api/admin/users', methods=['GET'])
@jwt_required()
@admin_required
def get_all_users():
    try:
        print("🔧 Admin Users API Called - Starting...")
        
        # Verify admin role first
        claims = get_jwt()
        print(f"🔐 JWT Claims: {claims}")
        
        if claims.get('role') != 'admin':
            print("❌ Not an admin user")
            return jsonify({"error": "Admin access required"}), 403

        # Get all users
        print("🔍 Querying users from database...")
        users = User.query.order_by(User.created_at.desc()).all()
        print(f"📊 Found {len(users)} users in database")

        users_data = []
        for user in users:
            try:
                # Get user stats
                user_notes = Note.query.filter_by(user_id=user.id).count()
                user_uploads = Analytics.query.filter_by(user_id=user.id, event_type="file_upload_started").count()
                last_login = Analytics.query.filter_by(user_id=user.id, event_type="user_login").order_by(
                    Analytics.timestamp.desc()
                ).first()

                # Determine status based on activity
                thirty_days_ago = datetime.now(timezone.utc) - timedelta(days=30)
                is_active = last_login and last_login.timestamp >= thirty_days_ago

                users_data.append({
                    "id": user.id,
                    "email": user.email,
                    "name": user.name or 'Anonymous',
                    "status": 'active' if is_active else 'inactive',
                    "created_at": user.created_at.isoformat(),
                    "last_login": last_login.timestamp.isoformat() if last_login else None,
                    "stats": {
                        "notes_count": user_notes,
                        "uploads_count": user_uploads
                    }
                })
            except Exception as user_error:
                print(f"❌ Error processing user {user.id}: {user_error}")
                # Add basic user info even if stats fail
                users_data.append({
                    "id": user.id,
                    "email": user.email,
                    "name": user.name or 'Anonymous',
                    "status": 'unknown',
                    "created_at": user.created_at.isoformat(),
                    "last_login": None,
                    "stats": {
                        "notes_count": 0,
                        "uploads_count": 0
                    }
                })

        print(f"📤 Returning {len(users_data)} users")
        return jsonify({"users": users_data}), 200
        
    except Exception as e:
        print(f"❌ Get all users error: {e}")
        import traceback
        print(f"📜 Traceback: {traceback.format_exc()}")
        return jsonify({"error": f"Failed to fetch users: {str(e)}"}), 500
@app.route('/api/admin/users/<user_id>/status', methods=['PATCH'])
@jwt_required()
@admin_required
def update_user_status(user_id):
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        data = request.get_json()
        new_status = data.get('status')

        if new_status not in ['active', 'inactive', 'suspended']:
            return jsonify({"error": "Invalid status"}), 400

        user = User.query.filter_by(id=user_id).first()
        if not user:
            return jsonify({"error": "User not found"}), 404

        # For now, we'll just track the status change in analytics
        # In a real app, you'd have a user status field in the User model
        track_event('admin_user_status_updated', {
            'user_id': user_id,
            'old_status': 'unknown',  # You'd need to track this properly
            'new_status': new_status,
            'admin_id': get_jwt_identity()
        })

        return jsonify({"message": f"User status updated to {new_status}"}), 200
    except Exception as e:
        print(f"Update user status error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to update user status"}), 500

@app.route('/api/admin/export/users.csv', methods=['GET'])
@jwt_required()
@admin_required
def export_users_csv():
    try:
        users = User.query.order_by(User.created_at.desc()).all()

        # Create CSV content
        csv_content = "ID,Name,Email,Status,Created At,Notes Count,Uploads Count,Tokens,XP,Level\n"

        for user in users:
            user_notes = Note.query.filter_by(user_id=user.id).count()
            user_uploads = Analytics.query.filter_by(user_id=user.id, event_type="file_upload_started").count()

            # Get user stats
            user_stats = get_stats(user.id)

            thirty_days_ago = datetime.now(timezone.utc) - timedelta(days=30)
            last_login = Analytics.query.filter_by(user_id=user.id, event_type="user_login").order_by(
                Analytics.timestamp.desc()
            ).first()
            is_active = last_login and last_login.timestamp >= thirty_days_ago
            status = 'active' if is_active else 'inactive'

            csv_content += f"{user.id},{user.name or 'Anonymous'},{user.email},{status},{user.created_at.isoformat()},{user_notes},{user_uploads},{user_stats.tokens if user_stats else 0},{user_stats.xp if user_stats else 0},{user_stats.level if user_stats else 1}\n"

        output = io.BytesIO()
        output.write(csv_content.encode('utf-8'))
        output.seek(0)

        return send_file(
            output,
            mimetype='text/csv',
            as_attachment=True,
            download_name=f'users_{datetime.now(pytz.timezone("Asia/Kolkata")).strftime("%Y%m%d_%H%M%S")}.csv'
        )
    except Exception as e:
        print(f"Export users CSV error: {e}")
        return jsonify({"error": "Failed to export users"}), 500

@app.route('/api/admin/stats', methods=['GET'])
@jwt_required()
@admin_required
def get_admin_stats():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        # User stats
        total_users = User.query.count()

        thirty_days_ago = datetime.now(timezone.utc) - timedelta(days=30)
        active_users = Analytics.query.filter(
            Analytics.event_type == 'user_login',
            Analytics.timestamp >= thirty_days_ago
        ).distinct(Analytics.user_id).count()

        inactive_users = total_users - active_users

        # Content stats
        total_notes = Note.query.count()
        total_uploads = Analytics.query.filter_by(event_type="file_upload_started").count()

        return jsonify({
            "total_users": total_users,
            "active_users": active_users,
            "inactive_users": inactive_users,
            "total_notes": total_notes,
            "total_uploads": total_uploads
        }), 200
    except Exception as e:
        print(f"Get admin stats error: {e}")
        return jsonify({"error": "Failed to fetch admin stats"}), 500

# ==================== ADMIN COMMUNITY MANAGEMENT ROUTES ====================

@app.route('/api/admin/community/posts', methods=['GET'])
@jwt_required()
@admin_required
def get_all_community_posts():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        posts = CommunityPost.query.order_by(CommunityPost.created_at.desc()).all()

        posts_data = []
        for post in posts:
            likes_count = CommunityLike.query.filter_by(post_id=post.id).count()

            # Determine status (for now, all posts are active unless flagged)
            flagged_events = Analytics.query.filter_by(
                event_type='admin_post_flagged',
                event_data=post.id
            ).count()
            status = 'flagged' if flagged_events > 0 else 'active'

            posts_data.append({
                "id": post.id,
                "content": post.content,
                "author_name": post.author_name,
                "author_email": post.author_email,
                "created_at": post.created_at.isoformat(),
                "likes_count": likes_count,
                "status": status
            })

        return jsonify({"posts": posts_data}), 200
    except Exception as e:
        print(f"Get all community posts error: {e}")
        return jsonify({"error": "Failed to fetch community posts"}), 500

@app.route('/api/admin/community/posts/<post_id>', methods=['DELETE'])
@jwt_required()
@admin_required
def delete_community_post(post_id):
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        post = CommunityPost.query.filter_by(id=post_id).first()
        if not post:
            return jsonify({"error": "Post not found"}), 404

        # Delete associated likes first
        CommunityLike.query.filter_by(post_id=post_id).delete()

        # Delete the post
        db.session.delete(post)
        db.session.commit()

        track_event('admin_post_deleted', {
            'post_id': post_id,
            'admin_id': get_jwt_identity()
        })

        return jsonify({"message": "Post deleted successfully"}), 200
    except Exception as e:
        print(f"Delete community post error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to delete post"}), 500

@app.route('/api/admin/community/posts/<post_id>/flag', methods=['PATCH'])
@jwt_required()
@admin_required
def flag_community_post(post_id):
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        post = CommunityPost.query.filter_by(id=post_id).first()
        if not post:
            return jsonify({"error": "Post not found"}), 404

        track_event('admin_post_flagged', {
            'post_id': post_id,
            'admin_id': get_jwt_identity()
        })

        return jsonify({"message": "Post flagged for review"}), 200
    except Exception as e:
        print(f"Flag community post error: {e}")
        return jsonify({"error": "Failed to flag post"}), 500

@app.route('/api/admin/community/stats', methods=['GET'])
@jwt_required()
@admin_required
def get_admin_community_stats():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        total_posts = CommunityPost.query.count()
        total_likes = CommunityLike.query.count()

        # Get flagged posts count (posts that have been flagged by admin)
        flagged_posts = Analytics.query.filter_by(event_type='admin_post_flagged').distinct(
            Analytics.event_data
        ).count()

        # Get active users (users who posted in last 30 days)
        thirty_days_ago = datetime.now(pytz.timezone('Asia/Kolkata')) - timedelta(days=30)
        active_users = CommunityPost.query.filter(
            CommunityPost.created_at >= thirty_days_ago
        ).distinct(CommunityPost.author_id).count()

        return jsonify({
            "total_posts": total_posts,
            "total_likes": total_likes,
            "flagged_posts": flagged_posts,
            "active_users": active_users
        }), 200
    except Exception as e:
        print(f"Get admin community stats error: {e}")
        return jsonify({"error": "Failed to fetch community stats"}), 500

# ==================== ADMIN SUPPORT MANAGEMENT ROUTES ====================

@app.route('/api/admin/support/tickets', methods=['GET'])
@jwt_required()
def get_all_support_tickets():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        tickets = SupportTicket.query.order_by(SupportTicket.created_at.desc()).all()

        tickets_data = []
        for ticket in tickets:
            tickets_data.append({
                "id": ticket.id,
                "title": ticket.title,
                "description": ticket.description,
                "category": ticket.category,
                "priority": ticket.priority,
                "status": ticket.status,
                "author_id": ticket.author_id,
                "created_at": ticket.created_at.isoformat(),
                "updated_at": ticket.updated_at.isoformat()
            })

        return jsonify({"tickets": tickets_data}), 200
    except Exception as e:
        print(f"Get all support tickets error: {e}")
        return jsonify({"error": "Failed to fetch support tickets"}), 500


@app.route('/api/admin/support/tickets/<ticket_id>', methods=['DELETE'])
@jwt_required()
@admin_required
def delete_support_ticket(ticket_id):
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        ticket = SupportTicket.query.filter_by(id=ticket_id).first()
        if not ticket:
            return jsonify({"error": "Ticket not found"}), 404

        db.session.delete(ticket)
        db.session.commit()

        track_event('admin_ticket_deleted', {
            'ticket_id': ticket_id,
            'admin_id': get_jwt_identity()
        })

        return jsonify({"message": "Ticket deleted successfully"}), 200
    except Exception as e:
        print(f"Delete support ticket error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to delete ticket"}), 500

@app.route('/api/admin/support/stats', methods=['GET'])
@jwt_required()
def get_admin_support_stats():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        total_tickets = SupportTicket.query.count()
        open_tickets = SupportTicket.query.filter_by(status='open').count()
        in_progress_tickets = SupportTicket.query.filter_by(status='in_progress').count()
        resolved_tickets = SupportTicket.query.filter_by(status='resolved').count()
        urgent_tickets = SupportTicket.query.filter_by(priority='urgent').count()

        return jsonify({
            "total_tickets": total_tickets,
            "open_tickets": open_tickets,
            "in_progress_tickets": in_progress_tickets,
            "resolved_tickets": resolved_tickets,
            "urgent_tickets": urgent_tickets
        }), 200
    except Exception as e:
        print(f"Get admin support stats error: {e}")
        return jsonify({"error": "Failed to fetch support stats"}), 500

# ==================== ADMIN ANALYTICS ROUTES ====================

@app.route('/api/admin/analytics', methods=['GET'])
@jwt_required()
def get_admin_analytics():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        time_range = request.args.get('range', '30d')

        # Calculate date range
        if time_range == '7d':
            days = 7
        elif time_range == '30d':
            days = 30
        elif time_range == '90d':
            days = 90
        elif time_range == '1y':
            days = 365
        else:
            days = 30

        start_date = datetime.now(pytz.timezone('Asia/Kolkata')) - timedelta(days=days)

        # User metrics
        total_users = User.query.count()
        new_users = User.query.filter(User.created_at >= start_date).count()

        # Active users (users who logged in within the time range)
        active_users = Analytics.query.filter(
            Analytics.event_type == 'user_login',
            Analytics.timestamp >= start_date
        ).distinct(Analytics.user_id).count()

        # Content metrics
        total_notes = Note.query.count()
        total_uploads = Analytics.query.filter(
            Analytics.event_type == 'file_upload_started',
            Analytics.timestamp >= start_date
        ).count()

        # AI usage
        ai_calls = Analytics.query.filter(
            Analytics.event_type == 'ai_generation',
            Analytics.timestamp >= start_date
        ).all()

        total_ai_calls = len(ai_calls)
        successful_ai_calls = len([c for c in ai_calls if c.event_data and c.event_data.get('success')])

        # Community metrics
        total_posts = CommunityPost.query.count()
        total_likes = CommunityLike.query.count()

        # Support metrics
        total_tickets = SupportTicket.query.count()
        resolved_tickets = SupportTicket.query.filter_by(status='resolved').count()

        # System health (mock data for now)
        avg_response_time = 150  # ms
        uptime_percentage = 99.5
        error_rate = 0.5

        return jsonify({
            "users": {
                "total": total_users,
                "active": active_users,
                "new": new_users,
                "churn_rate": 5.2  # Mock data
            },
            "content": {
                "total_notes": total_notes,
                "total_uploads": total_uploads,
                "avg_processing_time": 2.3,  # Mock data
                "ai_success_rate": round(successful_ai_calls / total_ai_calls * 100, 2) if total_ai_calls > 0 else 0
            },
            "community": {
                "total_posts": total_posts,
                "total_likes": total_likes,
                "active_users": active_users,
                "flagged_posts": 0  # Mock data
            },
            "support": {
                "total_tickets": total_tickets,
                "open_tickets": SupportTicket.query.filter_by(status='open').count(),
                "resolved_tickets": resolved_tickets,
                "avg_resolution_time": 24,  # Mock data in hours
                "satisfaction_rate": 85  # Mock data
            },
            "system": {
                "avg_response_time": avg_response_time,
                "uptime_percentage": uptime_percentage,
                "error_rate": error_rate
            }
        }), 200
    except Exception as e:
        print(f"Get admin analytics error: {e}")
        return jsonify({"error": "Failed to fetch analytics"}), 500

@app.route('/api/admin/analytics/export', methods=['GET'])
@jwt_required()
def export_admin_analytics():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        time_range = request.args.get('range', '30d')

        # Get analytics data
        analytics_data = get_admin_analytics()

        # Create CSV content
        csv_content = "Category,Metric,Value\n"

        # Users
        csv_content += f"Users,Total,{analytics_data.get_json()['users']['total']}\n"
        csv_content += f"Users,Active,{analytics_data.get_json()['users']['active']}\n"
        csv_content += f"Users,New,{analytics_data.get_json()['users']['new']}\n"
        csv_content += f"Users,Churn Rate,{analytics_data.get_json()['users']['churn_rate']}%\n"

        # Content
        csv_content += f"Content,Total Notes,{analytics_data.get_json()['content']['total_notes']}\n"
        csv_content += f"Content,Total Uploads,{analytics_data.get_json()['content']['total_uploads']}\n"
        csv_content += f"Content,Avg Processing Time,{analytics_data.get_json()['content']['avg_processing_time']}s\n"
        csv_content += f"Content,AI Success Rate,{analytics_data.get_json()['content']['ai_success_rate']}%\n"

        # Community
        csv_content += f"Community,Total Posts,{analytics_data.get_json()['community']['total_posts']}\n"
        csv_content += f"Community,Total Likes,{analytics_data.get_json()['community']['total_likes']}\n"
        csv_content += f"Community,Active Users,{analytics_data.get_json()['community']['active_users']}\n"
        csv_content += f"Community,Flagged Posts,{analytics_data.get_json()['community']['flagged_posts']}\n"

        # Support
        csv_content += f"Support,Total Tickets,{analytics_data.get_json()['support']['total_tickets']}\n"
        csv_content += f"Support,Open Tickets,{analytics_data.get_json()['support']['open_tickets']}\n"
        csv_content += f"Support,Resolved Tickets,{analytics_data.get_json()['support']['resolved_tickets']}\n"
        csv_content += f"Support,Avg Resolution Time,{analytics_data.get_json()['support']['avg_resolution_time']}h\n"
        csv_content += f"Support,Satisfaction Rate,{analytics_data.get_json()['support']['satisfaction_rate']}%\n"

        # System
        csv_content += f"System,Avg Response Time,{analytics_data.get_json()['system']['avg_response_time']}ms\n"
        csv_content += f"System,Uptime,{analytics_data.get_json()['system']['uptime_percentage']}%\n"
        csv_content += f"System,Error Rate,{analytics_data.get_json()['system']['error_rate']}%\n"

        output = io.BytesIO()
        output.write(csv_content.encode('utf-8'))
        output.seek(0)

        return send_file(
            output,
            mimetype='text/csv',
            as_attachment=True,
            download_name=f'analytics_{time_range}_{datetime.now(pytz.timezone("Asia/Kolkata")).strftime("%Y%m%d_%H%M%S")}.csv'
        )
    except Exception as e:
        print(f"Export admin analytics error: {e}")
        return jsonify({"error": "Failed to export analytics"}), 500

# ==================== ADMIN SETTINGS ROUTES ====================

@app.route('/api/admin/settings', methods=['GET'])
@jwt_required()
def get_admin_settings():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        # Get global settings from database
        global_settings = GlobalSettings.query.first()

        # Return current settings
        settings = {
            "global": {
                "free_uploads_per_day": global_settings.free_uploads_per_day if global_settings else 5,
                "free_chats_per_day": global_settings.free_chats_per_day if global_settings else 10,
                "daily_token_cap": global_settings.daily_token_cap if global_settings else 200,
                "streak_bonus_base": global_settings.streak_bonus_base if global_settings else 10
            },
            "system": {
                "maintenance_mode": False,
                "max_file_size": 50,
                "rate_limit_requests": 100,
                "rate_limit_window": 60
            },
            "ai": {
                "openai_enabled": bool(os.environ.get("OPENAI_API_KEY")),
                "gemini_enabled": False,
                "ollama_enabled": bool(os.environ.get("OLLAMA_URL")),
                "openai_model": os.environ.get("OPENAI_MODEL", "gpt-4o-mini"),
                "ollama_model": os.environ.get("OLLAMA_MODEL", "llama3:8b"),
                "max_tokens": 4000,
                "temperature": 0.7
            },
            "email": {
                "smtp_enabled": False,
                "smtp_host": "",
                "smtp_port": 587,
                "smtp_username": "",
                "smtp_password": "",
                "from_email": "",
                "from_name": ""
            },
            "security": {
                "jwt_expiry_days": 7,
                "password_min_length": 8,
                "enable_2fa": False,
                "session_timeout": 3600
            }
        }

        return jsonify(settings), 200
    except Exception as e:
        print(f"Get admin settings error: {e}")
        return jsonify({"error": "Failed to fetch settings"}), 500

@app.route('/api/admin/settings', methods=['PUT'])
@jwt_required()
def update_admin_settings():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        data = request.get_json()

        # In a real application, you'd validate and save these settings to a database
        # For now, we'll just acknowledge the update
        track_event('admin_settings_updated', {
            'admin_id': get_jwt_identity(),
            'settings_categories': list(data.keys())
        })

        return jsonify({"message": "Settings updated successfully"}), 200
    except Exception as e:
        print(f"Update admin settings error: {e}")
        return jsonify({"error": "Failed to update settings"}), 500

@app.route('/api/admin/settings/reset', methods=['POST'])
@jwt_required()
def reset_admin_settings():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        # Reset to default settings
        track_event('admin_settings_reset', {
            'admin_id': get_jwt_identity()
        })

        return jsonify({"message": "Settings reset to defaults"}), 200
    except Exception as e:
        print(f"Reset admin settings error: {e}")
        return jsonify({"error": "Failed to reset settings"}), 500

@app.route('/api/admin/settings/test-email', methods=['POST'])
@jwt_required()
def test_email_settings():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        # In a real application, you'd send a test email
        # For now, just acknowledge
        track_event('admin_test_email_sent', {
            'admin_id': get_jwt_identity()
        })

        return jsonify({"message": "Test email sent successfully"}), 200
    except Exception as e:
        print(f"Test email error: {e}")
        return jsonify({"error": "Failed to send test email"}), 500

@app.route('/api/admin/settings/update', methods=['POST'])
@jwt_required()
@admin_required
def update_global_settings():
    """Update global application settings"""
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        data = request.get_json()
        settings = GlobalSettings.query.get(1)

        if not settings:
            return jsonify({"error": "Global settings not found"}), 404

        # Update settings with provided values
        if 'free_uploads_per_day' in data:
            settings.free_uploads_per_day = data['free_uploads_per_day']
        if 'free_chats_per_day' in data:
            settings.free_chats_per_day = data['free_chats_per_day']
        if 'daily_token_cap' in data:
            settings.daily_token_cap = data['daily_token_cap']
        if 'streak_bonus_base' in data:
            settings.streak_bonus_base = data['streak_bonus_base']

        db.session.commit()

        # Track settings update
        track_event('admin_settings_updated', {
            'admin_id': get_jwt_identity(),
            'updated_fields': list(data.keys())
        })

        return jsonify({
            "success": True,
            "message": "Global settings updated successfully"
        }), 200

    except Exception as e:
        print(f"Update global settings error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to update settings"}), 500

# ==================== ADMIN SUPPORT ROUTES ====================

@app.route('/api/admin/support/tickets', methods=['GET'])
@jwt_required()
@admin_required
@track_usage
def admin_get_all_support_tickets():
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        tickets = SupportTicket.query.order_by(SupportTicket.created_at.desc()).all()

        tickets_data = []
        for ticket in tickets:
            tickets_data.append({
                "id": ticket.id,
                "title": ticket.title,
                "description": ticket.description,
                "category": ticket.category,
                "priority": ticket.priority,
                "status": ticket.status,
                "author_id": ticket.author_id,
                "created_at": ticket.created_at.isoformat(),
                "updated_at": ticket.updated_at.isoformat()
            })

        return jsonify({"tickets": tickets_data}), 200
    except Exception as e:
        print(f"Admin get all support tickets error: {e}")
        return jsonify({"error": "Failed to fetch support tickets"}), 500

@app.route('/api/admin/support/tickets/<ticket_id>/status', methods=['PATCH'])
@jwt_required()
@admin_required
@track_usage
def admin_update_support_ticket_status(ticket_id):
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        data = request.get_json()

        new_status = data.get('status')
        valid_statuses = ['open', 'in_progress', 'resolved', 'closed']

        if new_status not in valid_statuses:
            return jsonify({"error": "Invalid status"}), 400

        ticket = SupportTicket.query.filter_by(id=ticket_id).first()
        if not ticket:
            return jsonify({"error": "Ticket not found"}), 404

        old_status = ticket.status
        ticket.status = new_status
        ticket.updated_at = datetime.now(pytz.timezone('Asia/Kolkata'))

        db.session.commit()

        # Track admin status update
        track_event('admin_support_ticket_status_updated', {
            'ticket_id': ticket_id,
            'old_status': old_status,
            'new_status': new_status
        })

        return jsonify({
            "message": f"Ticket status updated to {new_status}",
            "ticket": {
                "id": ticket.id,
                "title": ticket.title,
                "status": ticket.status,
                "updated_at": ticket.updated_at.isoformat()
            }
        }), 200
    except Exception as e:
        print(f"Admin update support ticket status error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to update ticket status"}), 500
# ==================== ANALYTICS ROUTES ====================

@app.route('/api/analytics/stats', methods=['GET'])
@jwt_required()
@track_usage
def get_analytics_stats():
    try:
        user_id = get_jwt_identity()

        # Get user-specific stats
        total_notes = Note.query.filter_by(user_id=user_id).count()
        total_uploads = Analytics.query.filter_by(user_id=user_id, event_type="file_upload_started").count()
        total_exports = Analytics.query.filter_by(user_id=user_id, event_type="note_exported").count()

        # AI usage stats
        ai_calls = Analytics.query.filter_by(user_id=user_id, event_type="ai_generation").all()

        total_ai_calls = len(ai_calls)
        successful_ai_calls = len([c for c in ai_calls if c.event_data and c.event_data.get('success')])

        return jsonify({
            "total_notes": total_notes,
            "total_uploads": total_uploads,
            "total_exports": total_exports,
            "ai_usage": {
                "total_calls": total_ai_calls,
                "successful_calls": successful_ai_calls,
                "success_rate": round(successful_ai_calls / total_ai_calls * 100, 2) if total_ai_calls > 0 else 0
            }
        }), 200
    except Exception as e:
        print(f"Analytics error: {e}")
        return jsonify({"error": "Failed to fetch analytics"}), 500

@app.route('/api/chat', methods=['POST'])
@jwt_required()
@track_usage
def chat_with_ai():
    """Chat with AI assistant using context from selected notes"""
    try:
        user_id = get_jwt_identity()

        # Check if user can chat (free tier limits and tokens)
        if not can_chat(user_id):
            return jsonify({
                "error": "no_tokens",
                "message": "Buy tokens or upgrade to Premium."
            }), 402

        data = request.get_json()

        message = data.get('message', '').strip()
        note_id = data.get('note_id')  # Optional context note

        if not message:
            return jsonify({"error": "Message is required"}), 400

        # Get context from note if provided
        context = ""
        if note_id:
            note = Note.query.filter_by(id=note_id, user_id=user_id).first()
            if note:
                # Get relevant chunks using RAG
                context_chunks = retrieve_relevant_chunks(message, note_id, top_k=3)
                context = "\n\n".join(context_chunks) if context_chunks else note.content[:2000]
            else:
                return jsonify({"error": "Selected note not found"}), 404

        # Build prompt
        system_prompt = "You are an expert AI study assistant. Help students understand concepts, answer questions, and provide study guidance."

        if context:
            user_prompt = f"""
Based on the following study material:

{context}

Student's question: {message}

Please provide a helpful, accurate response that addresses their question using the context when relevant.
"""
        else:
            user_prompt = f"Student's question: {message}\n\nPlease provide a helpful response."

        # Check if user is premium for model selection
        subscription = get_subscription_status(user_id)
        is_premium = subscription and subscription['tier'] == 'premium'

        # Try AI providers
        response = None

        if openai_client:
            try:
                llm_client = LLMClient(provider="openai", use_premium=is_premium)
                response = llm_client.chat(system_prompt, user_prompt)
            except Exception as e:
                print(f"OpenAI chat failed: {e}")

        if response is None:
            try:
                llm_client = LLMClient(provider="ollama", use_premium=is_premium)
                response = llm_client.chat(system_prompt, user_prompt)
            except Exception as e:
                print(f"Ollama chat failed: {e}")

        if not response:
            return jsonify({"error": "AI service temporarily unavailable"}), 503

        # Update user streak and XP for chat activity
        update_streak_and_xp(user_id, "chat")

        # Log the chat interaction
        try:
            chat_log = ChatLog(
                id=str(uuid.uuid4()),
                user_id=user_id,
                message=message,
                response=response,
                tokens_used=3  # Approximate token cost
            )
            db.session.add(chat_log)
            db.session.commit()
        except Exception as log_error:
            print(f"Failed to log chat: {log_error}")
            # Don't fail the request if logging fails

        return jsonify({
            "response": response,
            "context_used": bool(note_id and context)
        }), 200

    except Exception as e:
        print(f"Chat error: {e}")
        return jsonify({"error": "Failed to process chat request"}), 500

@app.route("/api/chat", methods=["POST"])
@jwt_required()
def file_chat():
    try:
        user_id = get_jwt_identity()

        # Check if user can chat (free tier limits and tokens)
        if not can_chat(user_id):
            return jsonify({
                "error": "no_tokens",
                "message": "Buy tokens or upgrade to Premium."
            }), 402

        data = request.get_json() or {}
        note_id = data.get("note_id")
        user_message = data.get("message", "").strip()

        if not user_message:
            return jsonify({"error": "message required"}), 400

        # If no note/document selected
        if not note_id:
            return jsonify({
                "response": "Please select one of your generated notes/question papers before asking questions. "
                            "I can only help when a document is selected.",
                "context_used": False,
                "used_chunks": []
            }), 400

        # ---- Hybrid RAG retrieval ----
        relevant_chunks = retrieve_relevant_chunks(
            user_message, document_id=note_id, top_k=5
        )

        # ---- Natural-language system prompt (NO citations) ----
        system_prompt = (
            "You are a friendly, clear, and helpful study assistant. "
            "You must answer using ONLY the content found in the provided document excerpts. "
            "Do not mention chunks, citations, or sources. "
            "If the answer is not present in the document, say: "
            "'I couldn't find this information in your document.' "
            "Respond in a natural, conversational way.\n\n"
            "Here are the document excerpts:\n\n"
        )

        for c in relevant_chunks:
            system_prompt += c + "\n\n"

        # Check if user is premium for model selection
        subscription = get_subscription_status(user_id)
        is_premium = subscription and subscription['tier'] == 'premium'
        model = "gpt-4o" if is_premium else "gpt-4o-mini"

        # ---- OpenAI chat ----
        chat_resp = openai_client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            temperature=0.2
        )

        final_answer = chat_resp.choices[0].message["content"]

        # ---- Short preview for your context UI ----
        chunks_preview = [
            {"index": i, "preview": c[:200] + ("..." if len(c) > 200 else "")}
            for i, c in enumerate(relevant_chunks)
        ]

        # Update user streak and XP for file chat activity
        update_streak_and_xp(user_id, "file_chat")

        # Log the chat interaction
        try:
            chat_log = ChatLog(
                id=str(uuid.uuid4()),
                user_id=user_id,
                message=user_message,
                response=final_answer,
                tokens_used=3  # Approximate token cost
            )
            db.session.add(chat_log)
            db.session.commit()
        except Exception as log_error:
            print(f"Failed to log chat: {log_error}")
            # Don't fail the request if logging fails

        return jsonify({
            "response": final_answer,
            "context_used": True,
            "used_chunks": chunks_preview
        }), 200

    except Exception as e:
        print("Chat error:", e)
        return jsonify({"error": "internal_error"}), 500

@app.route('/api/notes/generate', methods=['POST'])
@jwt_required()
@track_usage
def generate_custom_notes():
    try:
        user_id = get_jwt_identity()
        data = request.get_json()
        document_id = data.get("document_id")
        topic = data.get("topic", "study notes")

        # Load text for this document (from your DB notes table)
        note = Note.query.filter_by(id=document_id, user_id=user_id).first()
        if not note:
            return jsonify({"error": "Document not found"}), 404

        # Chunk & embed if first time
        chunks = text_splitter.split_text(note.content)
        if chunks:
            try:
                chunk_and_embed_text(note.content, document_id, note.original_filename)
            except Exception as e:
                print(f"Embedding failed: {e}")
                # Continue without RAG if embedding fails

        # Use RAG to get relevant context
        context_chunks = retrieve_relevant_chunks(topic, document_id, top_k=5)
        context = "\n\n".join(context_chunks) if context_chunks else note.content[:4000]

        prompt = f"""
        You are a professional note generator for students.
        Create high-quality notes in bullet form.

        Topic: {topic}
        Context:
        {context}

        Provide summary, key points, formulas & tips.
        """

        # Check if user is premium for model selection
        subscription = get_subscription_status(user_id)
        is_premium = subscription and subscription['tier'] == 'premium'

        llm_client = LLMClient(provider="openai" if openai_client else "ollama", use_premium=is_premium)

        try:
            result = llm_client.chat(system_message="You are an expert study assistant.", user=prompt)
        except Exception as e:
            print(f"LLM failed: {e}")
            result = "AI temporarily unavailable. Try again later."

        return jsonify({
            "topic": topic,
            "notes": result
        }), 200
    except Exception as e:
        print(f"Generate notes error: {e}")
        return jsonify({"error": "Failed to generate notes"}), 500

@app.route('/api/papers/analyze', methods=['POST'])
@jwt_required()
@track_usage
def analyze_question_paper():
    try:
        user_id = get_jwt_identity()
        data = request.get_json()
        document_id = data.get("document_id")
        export_format = data.get("format", "json")

        doc = Note.query.filter_by(id=document_id, user_id=user_id).first()
        if not doc:
            return jsonify({"error": "Document not found"}), 404

        text = doc.content
        questions = [q for q in text.split("\n") if q.strip().endswith("?")]

        results = []
        for q in questions[:10]:  # Limit to first 10 questions
            prompt = f"""
            Analyze this exam question:

            "{q}"

            Classify into:
            - Topic
            - Difficulty (Easy/Medium/Hard)
            - One-line solution/hint

            Return format:
            Topic: <text>
            Difficulty: <text>
            Hint: <text>
            """

            # Check if user is premium for model selection
            subscription = get_subscription_status(user_id)
            is_premium = subscription and subscription['tier'] == 'premium'

            # Try OpenAI first, fallback to Ollama
            try:
                if openai_client:
                    llm_client = LLMClient(provider="openai", use_premium=is_premium)
                else:
                    llm_client = LLMClient(provider="ollama", use_premium=is_premium)
            except:
                llm_client = LLMClient(provider="ollama", use_premium=is_premium)

            try:
                analysis = llm_client.chat(system_message="You are an expert tutor.", user=prompt)
            except Exception as e:
                analysis = "Analysis unavailable"

            results.append({"question": q, "analysis": analysis})

        if export_format == "pdf":
            # PDF export disabled on shared hosting
            return jsonify({"error": "PDF export is not available on this server. Please use TXT export instead."}), 400

            track_event('paper_analysis_exported', {
                'document_id': document_id,
                'format': 'pdf'
            })

            return send_file(
                output,
                mimetype='application/pdf',
                as_attachment=True,
                download_name=filename
            )
        else:
            # Return JSON
            return jsonify(results), 200
    except Exception as e:
        print(f"Analyze paper error: {e}")
        return jsonify({"error": "Failed to analyze paper"}), 500

# ==================== USER PROFILE ROUTES ====================

@app.route('/api/user/profile', methods=['GET'])
@jwt_required()
@track_usage
def get_user_profile():
    """Get current user's profile information"""
    try:
        user_id = get_jwt_identity()
        user = User.query.filter_by(id=user_id).first()

        if not user:
            return jsonify({"error": "User not found"}), 404

        # Get user statistics
        total_notes = Note.query.filter_by(user_id=user_id).count()
        total_flashcards = Flashcard.query.filter_by(user_id=user_id).count()
        total_uploads = Analytics.query.filter_by(user_id=user_id, event_type="file_upload_started").count()

        # Get user stats (tokens, etc.)
        user_stats = get_stats(user_id)

        # Get referral information
        referral = Referral.query.filter_by(referrer_id=user_id).first()
        referral_code = referral.code if referral else None

        # Count successful referrals
        successful_referrals = Referral.query.filter_by(
            referrer_id=user_id,
            bonus_given=True
        ).count()

        # Get last login
        last_login = Analytics.query.filter_by(user_id=user_id, event_type="user_login").order_by(
            Analytics.timestamp.desc()
        ).first()

        return jsonify({
            "user": {
                "id": user.id,
                "email": user.email,
                "name": user.name or '',
                "created_at": user.created_at.isoformat(),
                "last_login": last_login.timestamp.isoformat() if last_login else None
            },
            "stats": {
                "total_notes": total_notes,
                "total_flashcards": total_flashcards,
                "total_uploads": total_uploads,
                "tokens": user_stats.tokens if user_stats else 0,
                "level": user_stats.level if user_stats else 1,
                "xp": user_stats.xp if user_stats else 0,
                "streak": user_stats.streak if user_stats else 0
            },
            "referral": {
                "code": referral_code,
                "successful_referrals": successful_referrals
            }
        }), 200
    except Exception as e:
        print(f"Get user profile error: {e}")
        return jsonify({"error": "Failed to fetch profile"}), 500

@app.route('/api/user/profile', methods=['PUT'])
@jwt_required()
@track_usage
def update_user_profile():
    """Update current user's profile information"""
    try:
        user_id = get_jwt_identity()
        user = User.query.filter_by(id=user_id).first()

        if not user:
            return jsonify({"error": "User not found"}), 404

        data = request.get_json()

        # Validate input
        name = data.get('name', '').strip()
        email = data.get('email', '').strip().lower()

        if not name:
            return jsonify({"error": "Name is required"}), 400

        if not email:
            return jsonify({"error": "Email is required"}), 400

        # Check if email is already taken by another user
        existing_user = User.query.filter_by(email=email).first()
        if existing_user and existing_user.id != user_id:
            return jsonify({"error": "Email is already in use"}), 400

        # Update user
        user.name = name
        user.email = email
        user.updated_at = datetime.now(pytz.timezone('Asia/Kolkata'))

        db.session.commit()

        track_event('profile_updated', {
            'user_id': user_id,
            'fields_updated': ['name', 'email']
        })

        return jsonify({
            "message": "Profile updated successfully",
            "user": {
                "id": user.id,
                "email": user.email,
                "name": user.name,
                "created_at": user.created_at.isoformat()
            }
        }), 200
    except Exception as e:
        print(f"Update user profile error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to update profile"}), 500

@app.route('/api/user/password', methods=['PUT'])
@jwt_required()
@track_usage
def change_user_password():
    """Change current user's password"""
    try:
        user_id = get_jwt_identity()
        user = User.query.filter_by(id=user_id).first()

        if not user:
            return jsonify({"error": "User not found"}), 404

        data = request.get_json()

        current_password = data.get('current_password')
        new_password = data.get('new_password')
        confirm_password = data.get('confirm_password')

        # Validate input
        if not current_password or not new_password or not confirm_password:
            return jsonify({"error": "All password fields are required"}), 400

        if new_password != confirm_password:
            return jsonify({"error": "New passwords do not match"}), 400

        if len(new_password) < 8:
            return jsonify({"error": "New password must be at least 8 characters long"}), 400

        # Verify current password
        if not check_password_hash(user.password, current_password):
            return jsonify({"error": "Current password is incorrect"}), 400

        # Update password
        user.password = generate_password_hash(new_password)
        user.updated_at = datetime.now(pytz.timezone('Asia/Kolkata'))

        db.session.commit()

        track_event('password_changed', {
            'user_id': user_id
        })

        return jsonify({"message": "Password changed successfully"}), 200
    except Exception as e:
        print(f"Change password error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to change password"}), 500

# ==================== USER NOTIFICATIONS ROUTES ====================
@app.route('/api/notifications', methods=['GET'])
@jwt_required()
def get_notifications():
    try:
        user_id = get_jwt_identity()
        notifications = Notification.query.filter_by(user_id=user_id).all()

        result = []
        for n in notifications:
            # SAFE created_at
            try:
                created_at = n.created_at.isoformat()
            except Exception as e:
                log_notification_error("created_at ERROR: " + str(e))
                created_at = None

            # SAFE expires_at
            try:
                expires_at = n.expires_at.isoformat() if n.expires_at else None
            except Exception as e:
                log_notification_error("expires_at ERROR: " + str(e))
                expires_at = None

            result.append({
                "id": n.id,
                "title": n.title,
                "message": n.message,
                "type": n.type,
                "is_read": n.is_read,
                "created_at": created_at,
                "expires_at": expires_at
            })

        return jsonify({"notifications": result}), 200

    except Exception as e:
        log_notification_error(e)
        return jsonify({"error": "notification route crashed"}), 500

@app.route('/api/notifications/<notification_id>/read', methods=['PATCH'])
@jwt_required()
@track_usage
def mark_notification_read(notification_id):
    try:
        user_id = get_jwt_identity()
        notification = Notification.query.filter_by(id=notification_id, user_id=user_id).first()

        if not notification:
            return jsonify({"error": "Notification not found"}), 404

        notification.is_read = True
        db.session.commit()

        track_event('notification_read', {
            'notification_id': notification_id,
            'user_id': user_id
        })

        return jsonify({"message": "Notification marked as read"}), 200
    except Exception as e:
        print(f"Mark notification read error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to mark notification as read"}), 500

@app.route('/api/notifications/mark-all-read', methods=['PATCH'])
@jwt_required()
@track_usage
def mark_all_notifications_read():
    try:
        user_id = get_jwt_identity()
        Notification.query.filter_by(user_id=user_id, is_read=False).update({'is_read': True})
        db.session.commit()

        track_event('all_notifications_read', {
            'user_id': user_id
        })

        return jsonify({"message": "All notifications marked as read"}), 200
    except Exception as e:
        print(f"Mark all notifications read error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to mark all notifications as read"}), 500

# ==================== ADMIN NOTIFICATIONS ROUTES ====================

@app.route('/api/admin/notifications', methods=['GET'])
@jwt_required()
@admin_required
def get_admin_notifications():
    try:
        # Get all notifications with recipient counts
        notifications = Notification.query.order_by(Notification.created_at.desc()).all()

        notifications_data = []
        for notification in notifications:
            # Get recipient count
            if notification.is_broadcast:
                # For broadcast notifications, count total users
                recipient_count = User.query.count()
                target = 'all'
            else:
                # For specific user notifications, count is 1
                recipient_count = 1 if notification.user_id else 0
                target = 'specific'

            # Map backend type to frontend type
            type_mapping_reverse = {
                'info': 'announcement',
                'success': 'update',
                'warning': 'alert',
                'error': 'alert'
            }
            frontend_type = type_mapping_reverse.get(notification.type, 'announcement')

            notifications_data.append({
                "id": notification.id,
                "title": notification.title,
                "message": notification.message,
                "type": frontend_type,
                "target": target,
                "created_at": notification.created_at.isoformat(),
                "expires_at": notification.expires_at.isoformat() if notification.expires_at else None,
                "recipient_count": recipient_count
            })

        return jsonify({"notifications": notifications_data}), 200
    except Exception as e:
        print(f"Get admin notifications error: {e}")
        return jsonify({"error": "Failed to fetch notifications"}), 500

@app.route('/api/admin/notifications', methods=['POST'])
@jwt_required()
@admin_required
def create_notification():
    try:
        claims = get_jwt()
        admin_id = get_jwt_identity()
        data = request.get_json()

        # Validate required fields
        if not data.get('title') or not data.get('message'):
            return jsonify({"error": "Title and message are required"}), 400

        title = data.get('title')
        message = data.get('message')
        frontend_type = data.get('type', 'announcement')
        target = data.get('target', 'all')
        expires_at = data.get('expires_at')

        # Map frontend types to backend types
        type_mapping = {
            'announcement': 'info',
            'alert': 'warning',
            'update': 'success',
            'maintenance': 'warning'
        }
        notification_type = type_mapping.get(frontend_type, 'info')

        # Determine target users
        target_users = []
        if target == 'all':
            # Broadcast to all users
            target_users = User.query.all()
        elif target == 'active':
            # Send to active users (logged in within last 30 days)
            thirty_days_ago = datetime.now(pytz.timezone('Asia/Kolkata')) - timedelta(days=30)
            active_user_ids = db.session.query(Analytics.user_id).filter(
                Analytics.event_type == 'user_login',
                Analytics.timestamp >= thirty_days_ago
            ).distinct().subquery()
            target_users = User.query.filter(User.id.in_(active_user_ids)).all()
        elif target == 'new':
            # Send to new users (registered within last 7 days)
            seven_days_ago = datetime.now(pytz.timezone('Asia/Kolkata')) - timedelta(days=7)
            target_users = User.query.filter(User.created_at >= seven_days_ago).all()
        else:
            return jsonify({"error": "Invalid target"}), 400

        # Create notifications
        notifications_created = 0
        for user in target_users:
            notification = Notification(
                id=str(uuid.uuid4()),
                user_id=user.id,
                title=title,
                message=message,
                type=notification_type,
                is_broadcast=(target == 'all'),
                expires_at=datetime.fromisoformat(expires_at.replace('Z', '+00:00')).astimezone(pytz.timezone('Asia/Kolkata')) if expires_at else None,
                created_by=admin_id,
                created_at=datetime.now(pytz.timezone('Asia/Kolkata'))
            )
            db.session.add(notification)
            notifications_created += 1

        db.session.commit()

        # Track notification creation
        track_event('admin_notification_created', {
            'admin_id': admin_id,
            'target': target,
            'recipient_count': notifications_created,
            'notification_type': notification_type
        })

        return jsonify({
            "message": f"Notification sent to {notifications_created} user(s)",
            "notifications_created": notifications_created
        }), 201
    except Exception as e:
        print(f"Create notification error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to create notification"}), 500

@app.route('/api/admin/notifications/<notification_id>', methods=['DELETE'])
@jwt_required()
@admin_required
def delete_notification(notification_id):
    try:
        notification = Notification.query.filter_by(id=notification_id).first()
        if not notification:
            return jsonify({"error": "Notification not found"}), 404

        db.session.delete(notification)
        db.session.commit()

        track_event('admin_notification_deleted', {
            'notification_id': notification_id,
            'admin_id': get_jwt_identity()
        })

        return jsonify({"message": "Notification deleted successfully"}), 200
    except Exception as e:
        print(f"Delete notification error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to delete notification"}), 500

@app.route('/api/admin/notifications/stats', methods=['GET'])
@jwt_required()
@admin_required
def get_admin_notification_stats():
    """Get notification statistics for admin dashboard"""
    try:
        # Total notifications sent
        total_sent = Notification.query.count()

        # Active subscribers (total users)
        active_subscribers = User.query.count()

        # Delivered notifications (all notifications are considered delivered)
        delivered = total_sent

        # Pending notifications (none, since they're sent immediately)
        pending = 0

        return jsonify({
            "total_sent": total_sent,
            "delivered": delivered,
            "pending": pending,
            "active_subscribers": active_subscribers
        }), 200
    except Exception as e:
        print(f"Get admin notification stats error: {e}")
        return jsonify({"error": "Failed to fetch notification stats"}), 500

@app.route('/api/', methods=['GET'])
def health_check():
    return jsonify({"message": "Impify API is running"}), 200

# Note: Removed WsgiToAsgi conversion - using Flask directly with uvicorn

# ==================== ADMIN DASHBOARD DATA PROVIDERS ====================

def admin_required(fn):
    @wraps(fn)
    @jwt_required()
    def wrapper(*args, **kwargs):
        claims = get_jwt()
        if claims.get("role") != "admin":
            return jsonify({"error": "Admin access required"}), 403
        return fn(*args, **kwargs)
    return wrapper

@app.route('/api/admin/payments', methods=['GET'])
@jwt_required()
@admin_required
def get_admin_payments():
    """Get all payment records for admin dashboard"""
    try:
        # Query real payment data from purchase_history table
        payments = PurchaseHistory.query.order_by(PurchaseHistory.created_at.desc()).all()

        payments_data = []
        for payment in payments:
            # Get user details
            user = User.query.filter_by(id=payment.user_id).first()
            user_name = user.name if user else 'Unknown'
            user_email = user.email if user else 'Unknown'

            payments_data.append({
                "id": payment.id,
                "user_id": payment.user_id,
                "user_name": user_name,
                "user_email": user_email,
                "pack_id": payment.pack_id,
                "tokens_credited": payment.tokens_credited,
                "amount_paid_in_inr": payment.amount_paid_in_inr,
                "payment_id": payment.payment_id,
                "status": "completed",  # All records in purchase_history are completed
                "created_at": payment.created_at.isoformat() if payment.created_at else None
            })

        return jsonify({"payments": payments_data}), 200
    except Exception as e:
        print(f"Get admin payments error: {e}")
        return jsonify({"error": "Failed to fetch payments"}), 500

@app.route('/api/admin/export/payments.csv', methods=['GET'])
@jwt_required()
@admin_required
def export_payments_csv():
    """Export payments as CSV"""
    try:
        # Query real payment data
        payments = PurchaseHistory.query.order_by(PurchaseHistory.created_at.desc()).all()

        csv_content = "ID,User ID,User Name,User Email,Pack ID,Tokens Credited,Amount Paid (INR),Payment ID,Status,Created At\n"

        for payment in payments:
            # Get user details
            user = User.query.filter_by(id=payment.user_id).first()
            user_name = user.name if user else 'Unknown'
            user_email = user.email if user else 'Unknown'

            created_at = payment.created_at.strftime('%Y-%m-%d %H:%M:%S') if payment.created_at else ''

            csv_content += f"{payment.id},{payment.user_id},{user_name},{user_email},{payment.pack_id or ''},{payment.tokens_credited},{payment.amount_paid_in_inr},{payment.payment_id},completed,{created_at}\n"

        output = io.BytesIO()
        output.write(csv_content.encode('utf-8'))
        output.seek(0)

        return send_file(
            output,
            mimetype='text/csv',
            as_attachment=True,
            download_name=f'payments_{datetime.now(pytz.timezone("Asia/Kolkata")).strftime("%Y%m%d_%H%M%S")}.csv'
        )
    except Exception as e:
        print(f"Export payments CSV error: {e}")
        return jsonify({"error": "Failed to export payments"}), 500

@app.route('/api/token-packs', methods=['GET'])
@jwt_required()
def get_token_packs():
    """Get available token packs"""
    try:
        # Return hardcoded token packs - in real app, this could be from DB
        token_packs = [
            {
                "id": "small",
                "name": "100 Tokens",
                "tokens": 100,
                "price_in_inr": 19,
                "description": "Perfect for light usage"
            },
            {
                "id": "medium",
                "name": "300 Tokens",
                "tokens": 300,
                "price_in_inr": 39,
                "description": "Most popular choice"
            },
            {
                "id": "large",
                "name": "1000 Tokens",
                "tokens": 1000,
                "price_in_inr": 99,
                "description": "Best value for heavy users"
            }
        ]

        return jsonify({"token_packs": token_packs}), 200
    except Exception as e:
        print(f"Get token packs error: {e}")
        return jsonify({"error": "Failed to fetch token packs"}), 500

@app.route('/api/admin/stats', methods=['GET'], endpoint='api_get_admin_stats')
@admin_required
def api_get_admin_stats():
    # User stats
    total_users = User.query.count()

    thirty_days_ago = datetime.now(timezone.utc) - timedelta(days=30)
    active_users = Analytics.query.filter(
        Analytics.event_type == 'user_login',
        Analytics.timestamp >= thirty_days_ago
    ).distinct(Analytics.user_id).count()

    inactive_users = total_users - active_users

    # Content stats
    total_notes = Note.query.count()
    total_uploads = Analytics.query.filter_by(event_type="file_upload_started").count()

    # Revenue stats from real payment data
    total_revenue_inr = db.session.query(db.func.sum(PurchaseHistory.amount_paid_in_inr)).scalar() or 0
    active_premium_users = Subscription.query.filter_by(active=True, tier='premium').count()

    return jsonify({
        "total_users": total_users,
        "active_users": active_users,
        "inactive_users": inactive_users,
        "total_notes": total_notes,
        "total_uploads": total_uploads,
        "total_revenue_in_inr": total_revenue_inr,
        "active_premium_users": active_premium_users
    }), 200

@app.route('/api/admin/users', methods=['GET'], endpoint='api_get_all_users')
@admin_required
def api_get_all_users():
    users = User.query.order_by(User.created_at.desc()).all()

    users_data = []
    for user in users:
        try:
            # Get user stats
            user_notes = Note.query.filter_by(user_id=user.id).count()
            user_uploads = Analytics.query.filter_by(user_id=user.id, event_type="file_upload_started").count()

            # Get user stats (tokens, etc.)
            user_stats = get_stats(user.id)

            # Determine status based on activity
            thirty_days_ago = datetime.now(timezone.utc) - timedelta(days=30)
            last_login = Analytics.query.filter_by(user_id=user.id, event_type="user_login").order_by(
                Analytics.timestamp.desc()
            ).first()
            is_active = last_login and last_login.timestamp >= thirty_days_ago

            users_data.append({
                "id": user.id,
                "email": user.email,
                "name": user.name or 'Anonymous',
                "status": 'active' if is_active else 'inactive',
                "created_at": user.created_at.isoformat(),
                "last_login": last_login.timestamp.isoformat() if last_login else None,
                "tokens": user_stats.tokens if user_stats else 0,
                "xp": user_stats.xp if user_stats else 0,
                "level": user_stats.level if user_stats else 1,
                "stats": {
                    "notes_count": user_notes,
                    "uploads_count": user_uploads
                }
            })
        except Exception as user_error:
            print(f"Error processing user {user.id}: {user_error}")
            # Add basic user info even if stats fail
            users_data.append({
                "id": user.id,
                "email": user.email,
                "name": user.name or 'Anonymous',
                "status": 'unknown',
                "created_at": user.created_at.isoformat(),
                "last_login": None,
                "tokens": 0,
                "xp": 0,
                "level": 1,
                "stats": {
                    "notes_count": 0,
                    "uploads_count": 0
                }
            })

    return jsonify({"users": users_data}), 200

@app.route('/api/admin/community/posts', methods=['GET'], endpoint='api_get_all_community_posts')
@admin_required
def api_get_all_community_posts():
    posts = CommunityPost.query.all()
    return jsonify({
        "posts": [{
            "id": p.id,
            "content": p.content,
            "author": p.author_email,
            "created_at": p.created_at.isoformat()
        } for p in posts]
    }), 200

@app.route('/api/admin/community/stats', methods=['GET'], endpoint='api_admin_community_stats')
@admin_required
def api_admin_community_stats():
    return jsonify({
        "total_posts": CommunityPost.query.count(),
        "total_likes": CommunityLike.query.count()
    }), 200

@app.route('/api/admin/support/tickets', methods=['GET'], endpoint='api_admin_support_tickets')
@admin_required
def api_admin_support_tickets():
    tickets = SupportTicket.query.order_by(SupportTicket.created_at.desc()).all()
    return jsonify({
        "tickets": [{
            "id": t.id,
            "title": t.title,
            "description": t.description,
            "category": t.category,
            "priority": t.priority,
            "status": t.status,
            "author_id": t.author_id,
            "created_at": t.created_at.isoformat(),
            "updated_at": t.updated_at.isoformat()
        } for t in tickets]
    }), 200

@app.route('/api/admin/support/stats', methods=['GET'], endpoint='api_admin_support_stats')
@admin_required
def api_admin_support_stats():
    return jsonify({
        "total_tickets": SupportTicket.query.count(),
        "open_tickets": SupportTicket.query.filter_by(status='open').count(),
        "in_progress_tickets": SupportTicket.query.filter_by(status='in_progress').count(),
        "resolved_tickets": SupportTicket.query.filter_by(status='resolved').count(),
        "urgent_tickets": SupportTicket.query.filter_by(priority='urgent').count()
    }), 200

@app.route('/api/admin/support/tickets/<ticket_id>/status', methods=['PATCH'], endpoint='api_admin_update_ticket_status')
@admin_required
def api_admin_update_ticket_status(ticket_id):
    data = request.get_json()
    new_status = data.get('status')

    valid_statuses = ['open', 'in_progress', 'resolved', 'closed']
    if new_status not in valid_statuses:
        return jsonify({"error": "Invalid status"}), 400

    ticket = SupportTicket.query.filter_by(id=ticket_id).first()
    if not ticket:
        return jsonify({"error": "Ticket not found"}), 404

    old_status = ticket.status
    ticket.status = new_status
    ticket.updated_at = datetime.now(pytz.timezone('Asia/Kolkata'))

    db.session.commit()

    track_event('admin_support_ticket_status_updated', {
        'ticket_id': ticket_id,
        'old_status': old_status,
        'new_status': new_status
    })

    return jsonify({
        "message": f"Ticket status updated to {new_status}",
        "ticket": {
            "id": ticket.id,
            "title": ticket.title,
            "status": ticket.status,
            "updated_at": ticket.updated_at.isoformat()
        }
    }), 200

@app.route('/api/admin/support/tickets/<ticket_id>', methods=['DELETE'], endpoint='api_admin_delete_ticket')
@admin_required
def api_admin_delete_ticket(ticket_id):
    ticket = SupportTicket.query.filter_by(id=ticket_id).first()
    if not ticket:
        return jsonify({"error": "Ticket not found"}), 404

    db.session.delete(ticket)
    db.session.commit()

    track_event('admin_ticket_deleted', {
        'ticket_id': ticket_id
    })

    return jsonify({"message": "Ticket deleted successfully"}), 200

@app.route('/api/admin/community/posts/<post_id>', methods=['DELETE'], endpoint='api_admin_delete_post')
@admin_required
def api_admin_delete_post(post_id):
    post = CommunityPost.query.filter_by(id=post_id).first()
    if not post:
        return jsonify({"error": "Post not found"}), 404

    # Delete associated likes first
    CommunityLike.query.filter_by(post_id=post_id).delete()

    # Delete the post
    db.session.delete(post)
    db.session.commit()

    track_event('admin_post_deleted', {
        'post_id': post_id
    })

    return jsonify({"message": "Post deleted successfully"}), 200

@app.route('/api/admin/users/<user_id>/status', methods=['PATCH'], endpoint='api_admin_update_user_status')
@admin_required
def api_admin_update_user_status(user_id):
    data = request.get_json()
    new_status = data.get('status')

    if new_status not in ['active', 'inactive', 'suspended']:
        return jsonify({"error": "Invalid status"}), 400

    user = User.query.filter_by(id=user_id).first()
    if not user:
        return jsonify({"error": "User not found"}), 404

    # For now, we'll just track the status change in analytics
    track_event('admin_user_status_updated', {
        'user_id': user_id,
        'new_status': new_status
    })

    return jsonify({"message": f"User status updated to {new_status}"}), 200

@app.route('/api/admin/users/<user_id>/tokens', methods=['POST'])
@jwt_required()
@admin_required
def admin_add_user_tokens(user_id):
    """Admin endpoint to add tokens to a user"""
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        data = request.get_json()
        amount = data.get('amount', 0)

        if amount <= 0:
            return jsonify({"error": "Amount must be positive"}), 400

        # Check if user exists
        user = User.query.filter_by(id=user_id).first()
        if not user:
            return jsonify({"error": "User not found"}), 404

        # Add tokens
        success = add_tokens(user_id, amount)
        if not success:
            return jsonify({"error": "Failed to add tokens"}), 500

        # Track admin action
        track_event('admin_tokens_added', {
            'admin_id': get_jwt_identity(),
            'user_id': user_id,
            'amount': amount
        })

        return jsonify({
            "success": True,
            "message": f"Added {amount} tokens to user {user.email}"
        }), 200

    except Exception as e:
        print(f"Admin add tokens error: {e}")
        return jsonify({"error": "Failed to add tokens"}), 500

@app.route('/api/admin/subscriptions', methods=['POST'])
@jwt_required()
@admin_required
def admin_create_subscription():
    """Admin endpoint to create premium subscriptions"""
    try:
        # Verify admin role
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({"error": "Admin access required"}), 403

        data = request.get_json()
        user_id = data.get('user_id')
        tier = data.get('tier', 'premium')
        expires_at_str = data.get('expires_at')

        if not user_id:
            return jsonify({"error": "user_id is required"}), 400

        # Check if user exists
        user = User.query.filter_by(id=user_id).first()
        if not user:
            return jsonify({"error": "User not found"}), 404

        # Parse expiration date
        expires_at = None
        if expires_at_str:
            try:
                expires_at = datetime.fromisoformat(expires_at_str.replace('Z', '+00:00'))
            except:
                return jsonify({"error": "Invalid expires_at format"}), 400

        # Create subscription
        subscription = Subscription(
            id=str(uuid.uuid4()),
            user_id=user_id,
            tier=tier,
            active=True,
            expires_at=expires_at,
            created_at=datetime.now(pytz.timezone('Asia/Kolkata'))
        )

        db.session.add(subscription)
        db.session.commit()

        # Track admin action
        track_event('admin_subscription_created', {
            'admin_id': get_jwt_identity(),
            'user_id': user_id,
            'tier': tier,
            'expires_at': expires_at_str
        })

        return jsonify({
            "success": True,
            "message": f"Created {tier} subscription for user {user.email}",
            "subscription": {
                "id": subscription.id,
                "tier": subscription.tier,
                "active": subscription.active,
                "expires_at": subscription.expires_at.isoformat() if subscription.expires_at else None
            }
        }), 201

    except Exception as e:
        print(f"Admin create subscription error: {e}")
        db.session.rollback()
        return jsonify({"error": "Failed to create subscription"}), 500

# ==================== TOKEN PURCHASE ROUTES ====================

@app.route('/api/token/purchase/initiate', methods=['POST'])
@jwt_required()
@track_usage
def initiate_token_purchase():
    """Initiate token purchase process"""
    try:
        user_id = get_jwt_identity()
        data = request.get_json()

        pack_type = data.get('pack_type')  # 'small', 'medium', 'large'
        payment_method = data.get('payment_method', 'razorpay')  # Default to Razorpay

        # Define token packs
        packs = {
            'small': {'tokens': 100, 'price': 19, 'currency': 'INR'},
            'medium': {'tokens': 300, 'price': 39, 'currency': 'INR'},
            'large': {'tokens': 1000, 'price': 99, 'currency': 'INR'}
        }

        if pack_type not in packs:
            return jsonify({"error": "Invalid pack type"}), 400

        pack = packs[pack_type]

        # Create purchase record (in a real app, you'd store this in a purchases table)
        purchase_id = str(uuid.uuid4())

        # Track purchase initiation
        track_event('token_purchase_initiated', {
            'user_id': user_id,
            'pack_type': pack_type,
            'tokens': pack['tokens'],
            'price': pack['price'],
            'purchase_id': purchase_id
        })

        # In a real implementation, you'd integrate with Razorpay/Stripe here
        # For now, return mock payment details
        return jsonify({
            "purchase_id": purchase_id,
            "pack_type": pack_type,
            "tokens": pack['tokens'],
            "price": pack['price'],
            "currency": pack['currency'],
            "payment_method": payment_method,
            "status": "initiated",
            # Mock payment gateway response
            "payment_details": {
                "order_id": f"order_{purchase_id}",
                "amount": pack['price'] * 100,  # Razorpay expects amount in paisa
                "currency": pack['currency']
            }
        }), 200

    except Exception as e:
        print(f"Initiate token purchase error: {e}")
        return jsonify({"error": "Failed to initiate purchase"}), 500

@app.route('/api/token/purchase/success', methods=['POST'])
@jwt_required()
@track_usage
def complete_token_purchase():
    """Complete successful token purchase"""
    try:
        user_id = get_jwt_identity()
        data = request.get_json()

        purchase_id = data.get('purchase_id')
        payment_id = data.get('payment_id')  # From payment gateway

        if not purchase_id or not payment_id:
            return jsonify({"error": "Purchase ID and payment ID are required"}), 400

        # In a real app, you'd verify the payment with the gateway
        # For now, assume payment is successful and add tokens

        # Mock token amounts based on purchase_id (in real app, store this in DB)
        token_amounts = {
            'small': 100,
            'medium': 300,
            'large': 1000
        }

        # Extract pack type from purchase_id or data
        pack_type = data.get('pack_type', 'small')  # In real app, retrieve from DB
        tokens_to_add = token_amounts.get(pack_type, 100)

        # Add tokens to user
        success = add_tokens(user_id, tokens_to_add)
        if not success:
            return jsonify({"error": "Failed to add tokens"}), 500

        # Track successful purchase
        track_event('token_purchase_completed', {
            'user_id': user_id,
            'purchase_id': purchase_id,
            'payment_id': payment_id,
            'pack_type': pack_type,
            'tokens_added': tokens_to_add
        })

        return jsonify({
            "message": "Token purchase completed successfully",
            "tokens_added": tokens_to_add,
            "new_balance": get_stats(user_id).tokens if get_stats(user_id) else 0
        }), 200

    except Exception as e:
        print(f"Complete token purchase error: {e}")
        return jsonify({"error": "Failed to complete purchase"}), 500

@app.route('/api/token/purchase/failure', methods=['POST'])
@jwt_required()
@track_usage
def handle_token_purchase_failure():
    """Handle failed token purchase"""
    try:
        user_id = get_jwt_identity()
        data = request.get_json()

        purchase_id = data.get('purchase_id')
        error_reason = data.get('error_reason', 'Unknown error')

        # Track failed purchase
        track_event('token_purchase_failed', {
            'user_id': user_id,
            'purchase_id': purchase_id,
            'error_reason': error_reason
        })

        return jsonify({
            "message": "Purchase failure recorded",
            "error_reason": error_reason
        }), 200

    except Exception as e:
        print(f"Handle token purchase failure error: {e}")
        return jsonify({"error": "Failed to handle purchase failure"}), 500

@app.route('/api/user/token-info', methods=['GET'])
@jwt_required()
@track_usage
def get_user_token_info_endpoint():
    """Get comprehensive token information for user"""
    try:
        user_id = get_jwt_identity()

        token_info = get_user_token_info(user_id)

        if not token_info:
            return jsonify({"error": "Failed to get token info"}), 500

        return jsonify(token_info), 200

    except Exception as e:
        print(f"Get user token info error: {e}")
        return jsonify({"error": "Failed to get token information"}), 500

@app.route('/api/dashboard/stats', methods=['GET'])
@jwt_required()
@track_usage
def dashboard_stats():
    """Get comprehensive dashboard statistics for user"""
    try:
        user_id = get_jwt_identity()

        # Notes count
        notes_count = Note.query.filter_by(user_id=user_id).count()

        # Flashcards count
        flashcards_count = Flashcard.query.filter_by(user_id=user_id).count()

        # Uploads count (notes created)
        uploads_count = notes_count

        # Calculate accuracy based on flashcard performance
        accuracy = 0
        if flashcards_count > 0:
            # Get all flashcards for the user
            flashcards = Flashcard.query.filter_by(user_id=user_id).all()
            if flashcards:
                total_correct = sum(card.correct_count for card in flashcards)
                total_reviews = sum(card.review_count for card in flashcards)
                if total_reviews > 0:
                    accuracy = round((total_correct / total_reviews) * 100, 1)

        # Chat logs (for future analytics)
        chats_today = ChatLog.query.filter(
            ChatLog.user_id == user_id,
            ChatLog.timestamp >= date.today()
        ).count()

        # Get full token+xp+streak stats
        token_info = get_user_token_info(user_id)

        return jsonify({
            "notes": notes_count,
            "flashcards": flashcards_count,
            "uploads": uploads_count,
            "accuracy": accuracy,
            "chats_today": chats_today,
            "tokens": token_info['current_tokens'],
            "monthly_tokens_remaining": token_info['monthly_tokens_remaining'],
            "streak": token_info['streak_days'],
            "level": token_info['level'],
            "xp": token_info['xp'],
            "xp_to_next_level": token_info['xp_to_next_level'],
            "days_until_reset": token_info['days_until_reset'],
            "subscription": token_info['subscription']
        }), 200

    except Exception as e:
        print(f"Dashboard stats error: {e}")
        return jsonify({"error": "Failed to get dashboard stats"}), 500

@app.route('/api/user/activity', methods=['POST'])
@jwt_required()
@track_usage
def record_user_activity():
    """Record user activity for streak and XP tracking"""
    try:
        user_id = get_jwt_identity()
        data = request.get_json()
        activity_type = data.get('activity_type', 'login')

        success = update_streak_and_xp(user_id, activity_type)

        if success:
            return jsonify({"message": "Activity recorded successfully"}), 200
        else:
            return jsonify({"error": "Failed to record activity"}), 500

    except Exception as e:
        print(f"Record user activity error: {e}")
        return jsonify({"error": "Failed to record activity"}), 500

# Register admin subscription blueprint
from routes.admin_subscriptions import bp as admin_subscriptions_bp
app.register_blueprint(admin_subscriptions_bp)

# Temporary error handler for debugging (remove in production)
import traceback
from flask import jsonify

@app.errorhandler(Exception)
def handle_all_errors(e):
    tb = traceback.format_exc()
    print("ERROR:", str(e))
    print(tb)
    return jsonify({"error": str(e), "trace": tb}), 500

  # Production uses Passenger WSGI — do NOT run app here!